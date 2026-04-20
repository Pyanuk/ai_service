from __future__ import annotations

from copy import deepcopy
import json
from pathlib import Path
import zipfile

from docx import Document


def _paragraph_signature(paragraph) -> tuple:
    run = paragraph.runs[0] if paragraph.runs else None
    return (
        paragraph.style.name,
        paragraph.alignment,
        paragraph.paragraph_format.left_indent,
        paragraph.paragraph_format.first_line_indent,
        run.font.name if run else None,
        run.font.size if run else None,
        run.bold if run else None,
        run.italic if run else None,
        run.underline if run else None,
    )


def _find_paragraph_index(document: Document, text: str) -> int:
    return next(index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == text)


def test_resolve_standard_by_fgos_url(client):
    payload = {
        "fgos_url": "https://fgos.ru/fgos/fgos-09-02-07-informacionnye-sistemy-i-programmirovanie-1547/",
        "course_name": "Инженерия промптов для команд",
        "brief_description": "Практика автоматизации и прикладного программирования с ИИ.",
    }

    response = client.post("/api/standards/resolve", json=payload)

    assert response.status_code == 200
    body = response.json()
    assert body["supported"] is True
    assert body["fgos_code"] == "09.02.07"
    assert body["standard_profile_id"] == "fgos_spo_09_02_07"
    assert body["resolved_track_id"] == "programmer"
    assert any(track["track_id"] == "web" for track in body["supported_tracks"])


def test_resolve_standard_reports_unsupported_fgos(client):
    response = client.post("/api/standards/resolve", json={"fgos_code": "09.02.13"})

    assert response.status_code == 200
    body = response.json()
    assert body["supported"] is False
    assert body["fgos_code"] == "09.02.13"
    assert body["standard_profile_id"] is None


def test_resolve_standard_auto_registers_profile_from_url(client):
    response = client.post(
        "/api/standards/resolve",
        json={
            "fgos_url": "https://example.com/fgos-09-02-09.pdf",
            "course_name": "Backend-разработка на JavaScript (Node.js, Express)",
            "brief_description": "Курс по backend-разработке, Node.js, Express и REST API.",
        },
    )

    assert response.status_code == 200
    body = response.json()
    assert body["supported"] is True
    assert body["fgos_code"] == "09.02.09"
    assert body["standard_profile_id"] == "auto_fgos_spo_09_02_09"
    assert body["resolved_track_id"] == "generic"


def test_generate_draft_with_auto_registered_profile(client, seed_payload):
    resolved = client.post(
        "/api/standards/resolve",
        json={
            "fgos_url": "https://example.com/fgos-09-02-09.pdf",
            "course_name": "Backend-разработка на JavaScript (Node.js, Express)",
            "brief_description": "Курс по backend-разработке, Node.js, Express и REST API.",
        },
    ).json()

    payload = dict(seed_payload)
    payload["course_name"] = "Backend-разработка на JavaScript (Node.js, Express)"
    payload["brief_description"] = "Курс по backend-разработке, Node.js, Express и REST API."
    payload["professional_area"] = "Backend-разработка WEB-приложений и серверной логики."
    payload["training_goal"] = "Освоение backend-разработки на JavaScript, Node.js и Express."
    payload["constraints"] = dict(seed_payload["constraints"])
    payload["constraints"]["standard_profile_id"] = resolved["standard_profile_id"]
    payload["constraints"]["standard_track_id"] = resolved["resolved_track_id"]

    response = client.post("/api/course-drafts/generate", json=payload)

    assert response.status_code == 200
    body = response.json()
    assert body["draft"]["seed"]["constraints"]["standard_profile_id"] == "auto_fgos_spo_09_02_09"
    assert body["draft"]["seed"]["constraints"]["standard_track_id"] == "generic"


def test_resolve_pdf_dynamic_profile_uses_competencies_from_pdf(client, seed_payload, monkeypatch):
    service = client.app.state.services["standards_service"]
    monkeypatch.setattr(
        service,
        "extract_text_from_pdf_bytes",
        lambda _: "\n".join(
            [
                "ФГОС СПО 09.02.02 Компьютерные сети",
                "ПК 1.1 Выполнять проектирование кабельной структуры компьютерной сети.",
                "ПК 1.2 Осуществлять выбор технологии, инструментальных средств и средств вычислительной техники при организации процесса разработки и исследования объектов профессиональной деятельности.",
                "ПК 2.1 Администрировать локальные вычислительные сети и принимать меры по устранению возможных сбоев.",
                "ПК 2.2 Администрировать сетевые ресурсы в информационных системах.",
            ]
        ),
    )

    resolved = client.post(
        "/api/standards/resolve-pdf",
        data={"fgos_code": "09.02.02"},
        files={"fgos_pdf": ("fgos.pdf", b"%PDF-1.4 fake", "application/pdf")},
    )

    assert resolved.status_code == 200
    body = resolved.json()
    assert body["standard_profile_id"] == "auto_fgos_spo_09_02_02"
    assert body["detected_competencies"] == [
        "ПК 1.1 Выполнять проектирование кабельной структуры компьютерной сети.",
        "ПК 1.2 Осуществлять выбор технологии, инструментальных средств и средств вычислительной техники при организации процесса разработки и исследования объектов профессиональной деятельности.",
        "ПК 2.1 Администрировать локальные вычислительные сети и принимать меры по устранению возможных сбоев.",
        "ПК 2.2 Администрировать сетевые ресурсы в информационных системах.",
    ]

    payload = deepcopy(seed_payload)
    payload["course_name"] = "Администрирование и защита компьютерных сетей"
    payload["qualification"] = "специалист по компьютерным сетям"
    payload["professional_area"] = "Проектирование, администрирование и сопровождение компьютерных сетей"
    payload["training_goal"] = "Освоение проектирования и администрирования компьютерных сетей"
    payload["brief_description"] = "Программа по проектированию и администрированию компьютерных сетей."
    payload["constraints"]["standard_profile_id"] = body["standard_profile_id"]
    payload["constraints"]["standard_track_id"] = body["resolved_track_id"]

    generated = client.post("/api/course-drafts/generate", json=payload)
    assert generated.status_code == 200
    draft = generated.json()["draft"]

    competencies = draft["activity_matrix"][0]["competencies"]
    assert "ПК 1.1 Выполнять проектирование кабельной структуры компьютерной сети." in competencies
    assert "ПК 1.2 Осуществлять выбор технологии, инструментальных средств и средств вычислительной техники при организации процесса разработки и исследования объектов профессиональной деятельности." in competencies
    assert "ПК 2.1 Администрировать локальные вычислительные сети и принимать меры по устранению возможных сбоев." in competencies
    assert "ПК 2.2 Администрировать сетевые ресурсы в информационных системах." in competencies


def test_resolve_pdf_rejects_dynamic_profile_without_extracted_competencies(client, monkeypatch):
    service = client.app.state.services["standards_service"]
    monkeypatch.setattr(
        service,
        "extract_text_from_pdf_bytes",
        lambda _: "ФГОС СПО 09.02.13 без распознаваемого блока профессиональных компетенций",
    )

    resolved = client.post(
        "/api/standards/resolve-pdf",
        data={"fgos_code": "09.02.13"},
        files={"fgos_pdf": ("fgos.pdf", b"%PDF-1.4 fake", "application/pdf")},
    )

    assert resolved.status_code == 422
    assert "Не удалось извлечь профессиональные компетенции из PDF ФГОС 09.02.13" in resolved.json()["detail"]

def test_generate_draft_sanitizes_saved_dynamic_profile_competencies(client, seed_payload):
    settings = client.app.state.services["settings"]
    registry_path = settings.service_root / "storage" / "standards" / "profiles.json"
    registry_path.parent.mkdir(parents=True, exist_ok=True)
    registry_path.write_text(
        json.dumps(
            {
                "auto_fgos_spo_09_02_02": {
                    "profile_id": "auto_fgos_spo_09_02_02",
                    "fgos_code": "09.02.02",
                    "title": "ФГОС СПО 09.02.02",
                    "order_title": "приказ",
                    "source_url": "https://registry.local/fgos/09.02.02",
                    "professional_area": "Сетевые технологии",
                    "qualification_level": "Уровень определяется ФГОС 09.02.02.",
                    "parallel_education_note": "Допускается параллельное освоение.",
                    "audience_requirements": [
                        "Лица, имеющие среднее профессиональное и (или) высшее образование.",
                    ],
                    "additional_requirements": [
                        "Наличие базовых навыков работы с компьютером.",
                    ],
                    "entry_requirements": "Вступительные испытания не предусмотрены.",
                    "track_id": "generic",
                    "qualification_title": "специалист по компьютерным сетям",
                    "course_name": "Проектирование и администрирование защищенных компьютерных сетей",
                    "competencies": [
                        "ПК 1.1 Проектировать кабельную структуру компьютерной сети. 5.2.2. Организация сетевого администрирования.",
                        "ПК 1.2 Администрировать сетевые ресурсы. fgos.ru 20.04.2026.",
                        "ПК 2.1 Эксплуатировать объекты сетевой инфраструктуры. Приложение к ФГОС СПО 09.02.02.",
                        "ПК 2.2 Обеспечивать защиту информации в сети. ПК 2.3 Выполнять мониторинг сети и диагностику неисправностей.",
                    ],
                }
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    payload = deepcopy(seed_payload)
    payload["course_name"] = "Проектирование и администрирование защищенных компьютерных сетей"
    payload["qualification"] = "специалист по компьютерным сетям"
    payload["professional_area"] = "Проектирование, настройка, эксплуатация, мониторинг и защита компьютерных сетей"
    payload["training_goal"] = "Освоение проектирования, настройки и администрирования компьютерных сетей"
    payload["brief_description"] = "Программа посвящена проектированию и сопровождению компьютерных сетей."
    payload["constraints"]["standard_profile_id"] = "auto_fgos_spo_09_02_02"
    payload["constraints"]["standard_track_id"] = "generic"

    generated = client.post("/api/course-drafts/generate", json=payload)

    assert generated.status_code == 200
    draft = generated.json()["draft"]
    competencies = [item for function in draft["labor_functions"] for item in function["competencies"]]

    assert competencies == [
        "ПК 1.1 Проектировать кабельную структуру компьютерной сети.",
        "ПК 1.2 Администрировать сетевые ресурсы.",
        "ПК 2.1 Эксплуатировать объекты сетевой инфраструктуры.",
        "ПК 2.2 Обеспечивать защиту информации в сети.",
        "ПК 2.3 Выполнять мониторинг сети и диагностику неисправностей.",
    ]
    assert all("fgos.ru" not in item.lower() for item in competencies)
    assert all("5.2.2" not in item for item in competencies)
    assert all("Приложение" not in item for item in competencies)


def test_backend_course_uses_exact_site_topics_and_practices(client, seed_payload):
    resolved = client.post(
        "/api/standards/resolve",
        json={
            "fgos_url": "https://example.com/fgos-09-02-09.pdf",
            "course_name": "Backend-разработка на JavaScript (Node.js, Express)",
            "brief_description": "Курс по backend-разработке, Node.js, Express и REST API.",
        },
    ).json()

    payload = dict(seed_payload)
    payload["course_name"] = "Backend-разработка на JavaScript (Node.js, Express)"
    payload["hours"] = 256
    payload["qualification"] = "Разработчик WEB и мультимедийных приложений"
    payload["professional_area"] = "Backend-разработка WEB-приложений, создание серверной логики и REST API"
    payload["training_goal"] = "Освоение JavaScript, Node.js, Express, работы с PostgreSQL и MongoDB, разработки REST API и интеграции backend с клиентской частью."
    payload["brief_description"] = "Программа по JavaScript, Node.js, Express, PostgreSQL, MongoDB, JWT и REST API."
    payload["modules_seed"] = [
        {
            "name": "Программирование на языке JavaScript",
            "desired_hours": 64,
            "summary": "Основы JavaScript, DOM, события, асинхронный код, модули, localStorage/sessionStorage, работа с API и JSON.",
        },
        {
            "name": "Работа с базами данных и интеграция с backend",
            "desired_hours": 64,
            "summary": "Node.js, Express, middleware, JWT, CORS, работа с PostgreSQL и MongoDB, создание REST API, логирование и тестирование backend-приложений.",
        },
        {
            "name": "Аутентификация, авторизация и продвинутый backend",
            "desired_hours": 64,
            "summary": "SQL, PostgreSQL, MongoDB, индексация, CRUD, безопасность, резервное копирование, проектирование баз данных и многозвенные приложения.",
        },
        {
            "name": "Разработка REST API и интеграция с клиентской частью WEB-приложений",
            "desired_hours": 64,
            "summary": "REST API, Express, параметры запросов, JWT, CORS, роли, интеграция с фронтендом, AJAX, обработка ошибок, кэширование, версионирование, тестирование и документация API.",
        },
    ]
    payload["constraints"] = dict(seed_payload["constraints"])
    payload["constraints"]["standard_profile_id"] = resolved["standard_profile_id"]
    payload["constraints"]["standard_track_id"] = resolved["resolved_track_id"]
    payload["source_url"] = "https://25-12.ru/courses/backend-%D1%80%D0%B0%D0%B7%D1%80%D0%B0%D0%B1%D0%BE%D1%82%D0%BA%D0%B0-%D0%BD%D0%B0-javascript-node-js-express/"

    generated = client.post("/api/course-drafts/generate", json=payload)
    assert generated.status_code == 200
    draft = generated.json()["draft"]

    assert "Введение в JavaScript и настройка окружения" in draft["working_programs_block"]
    assert "Практическая работа №18. Запрос к API и отображение данных" in draft["working_programs_block"]
    assert "Введение в Node.js и настройка окружения" in draft["working_programs_block"]
    assert "Практическая работа №18. Написание тестов для маршрутов Express" in draft["working_programs_block"]
    assert "Практическая работа №18. Создание документации REST API" in draft["working_programs_block"]

    exported = client.post(f"/api/course-drafts/{generated.json()['draft_id']}/export-docx")
    assert exported.status_code == 200
    document = Document(Path(exported.json()["document_path"]))
    thematic_text = "\n".join(cell.text for row in document.tables[8].rows for cell in row.cells)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Введение в JavaScript и настройка окружения" in thematic_text
    assert "Создание документации для API" in thematic_text
    assert "Практическая работа №14. Реализация полнофункционального REST API" in paragraph_text
    assert "Практическая работа №18. Создание документации REST API" in paragraph_text


def test_source_url_outline_is_used_for_new_course_without_manual_mapping(client, seed_payload, monkeypatch):
    class FakeResponse:
        text = """
        <html><body>
        <h3>Содержание</h3>
        <div>Модуль 1. Работа с HTTP API</div>
        <div>1.1</div><div>Введение в HTTP API</div>
        <div>1.2</div><div>Практическая работа №1. Отправка GET-запроса</div>
        <div>1.3</div><div>Параметры и заголовки запросов</div>
        <div>1.4</div><div>Практическая работа №2. Работа с headers и query params</div>
        <div>Модуль 2. Интеграция сервисов</div>
        <div>2.1</div><div>Основы webhook</div>
        <div>2.2</div><div>Практическая работа №1. Настройка webhook</div>
        <div>2.3</div><div>Обработка ответов внешних сервисов</div>
        <div>2.4</div><div>Практическая работа №2. Интеграция двух API</div>
        </body></html>
        """

        def raise_for_status(self):
            return None

    monkeypatch.setattr("app.services.draft_builder.requests.get", lambda *args, **kwargs: FakeResponse())

    payload = dict(seed_payload)
    payload["course_name"] = "Работа с HTTP API и интеграцией сервисов"
    payload["brief_description"] = "Курс по API и интеграции сервисов."
    payload["source_url"] = "https://25-12.ru/courses/custom-api-course/"
    payload["modules_seed"] = [
        {
            "name": "Работа с HTTP API",
            "desired_hours": 24,
            "summary": "GET, POST, headers, query params.",
        },
        {
            "name": "Интеграция сервисов",
            "desired_hours": 24,
            "summary": "Webhook, обработка ответов и связка API.",
        },
    ]

    generated = client.post("/api/course-drafts/generate", json=payload)
    assert generated.status_code == 200
    draft = generated.json()["draft"]

    assert "Введение в HTTP API" in draft["working_programs_block"]
    assert "Практическая работа №1. Отправка GET-запроса" in draft["working_programs_block"]
    assert "Основы webhook" in draft["working_programs_block"]
    assert "Практическая работа №2. Интеграция двух API" in draft["working_programs_block"]

    exported = client.post(f"/api/course-drafts/{generated.json()['draft_id']}/export-docx")
    document = Document(Path(exported.json()["document_path"]))
    thematic_text = "\n".join(cell.text for row in document.tables[8].rows for cell in row.cells)
    assert "Введение в HTTP API" in thematic_text
    assert "Обработка ответов внешних сервисов" in thematic_text


def test_resolve_standard_by_pdf(client, monkeypatch):
    service = client.app.state.services["standards_service"]
    monkeypatch.setattr(
        service,
        "extract_text_from_pdf_bytes",
        lambda _: "ФГОС СПО 09.02.11 Разработка и управление программным обеспечением",
    )

    response = client.post(
        "/api/standards/resolve-pdf",
        data={"fgos_code": "09.02.11"},
        files={"fgos_pdf": ("fgos.pdf", b"%PDF-1.4 fake", "application/pdf")},
    )

    assert response.status_code == 200
    body = response.json()
    assert body["supported"] is True
    assert body["fgos_code"] == "09.02.11"
    assert body["standard_profile_id"] == "fgos_spo_09_02_11"
    assert body["resolved_track_id"] == "devops_infrastructure"


def test_generate_draft(client, seed_payload):
    response = client.post("/api/course-drafts/generate", json=seed_payload)
    assert response.status_code == 200
    body = response.json()
    draft = body["draft"]

    assert body["draft_id"]
    assert draft["program_card"]["course_name"] == seed_payload["course_name"]
    assert len(draft["modules"]) == 2
    assert draft["seed"]["constraints"]["standard_profile_id"] == "fgos_spo_09_02_07"
    assert draft["seed"]["constraints"]["standard_track_id"] == "programmer"
    assert "09.02.07" in draft["general_characteristics"]["standards_basis"]
    assert "09.12.2016" in draft["general_characteristics"]["standards_basis"]
    assert "17.12.2020" in draft["general_characteristics"]["standards_basis"]
    assert draft["labor_functions"][0]["code_level"].startswith("A/")
    assert "Трудовые функции:" in draft["activity_matrix"][0]["competencies"]


def test_generate_draft_rejects_invalid_payload(client, seed_payload):
    broken = dict(seed_payload)
    broken["modules_seed"] = []
    response = client.post("/api/course-drafts/generate", json=broken)
    assert response.status_code == 422


def test_generate_draft_rejects_unsupported_standard_profile(client, seed_payload):
    payload = dict(seed_payload)
    payload["constraints"] = dict(seed_payload["constraints"])
    payload["constraints"]["standard_profile_id"] = "unknown-profile"

    response = client.post("/api/course-drafts/generate", json=payload)

    assert response.status_code == 422
    assert "Неподдерживаемый профиль стандартов" in response.json()["detail"]


def test_generate_devops_draft_for_fgos_09_02_11(client, devops_seed_payload):
    response = client.post("/api/course-drafts/generate", json=devops_seed_payload)

    assert response.status_code == 200
    body = response.json()
    draft = body["draft"]

    assert draft["program_card"]["course_name"] == devops_seed_payload["course_name"]
    assert draft["program_card"]["hours"] == 256
    assert draft["seed"]["constraints"]["standard_profile_id"] == "fgos_spo_09_02_11"
    assert draft["seed"]["constraints"]["standard_track_id"] == "devops_infrastructure"
    assert "09.02.11" in draft["general_characteristics"]["program_goal"]
    assert "24.02.2025 № 138" in draft["general_characteristics"]["standards_basis"]
    assert "специалист по информационным системам" in draft["general_characteristics"]["final_attestation_result"]
    assert draft["labor_functions"][0]["code_level"].startswith("A/")
    assert "Трудовые функции:" in draft["activity_matrix"][0]["competencies"]
    assert "Docker" in draft["activity_matrix"][0]["practical_experience"]


def test_devops_thematic_plan_is_detailed_and_keeps_font(client, devops_seed_payload):
    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")

    assert response.status_code == 200

    document = Document(Path(response.json()["document_path"]))
    thematic_table = document.tables[8]

    assert len(thematic_table.rows) >= 45

    thematic_text = "\n".join(cell.text for row in thematic_table.rows for cell in row.cells)
    assert "Dockerfile" in thematic_text
    assert "CI/CD" in thematic_text
    assert "Python" in thematic_text

    first_theme_run = next(
        run
        for row in thematic_table.rows[5:]
        for cell in row.cells
        for paragraph in cell.paragraphs
        for run in paragraph.runs
        if run.text.strip()
    )
    assert first_theme_run.font.name == "Montserrat"
    assert first_theme_run.font.size.pt == 9


def test_devops_working_programs_match_course_outline(client, devops_seed_payload):
    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    draft = generated["draft"]

    block = draft["working_programs_block"]
    assert "Тема 1.1. Введение в Python и установка среды разработки" in block
    assert "Практическая работа №1. Установка Python и запуск первой программы" in block
    assert "Практическая работа №18. Поиск и замена данных с использованием регулярных выражений." in block
    assert "Практическая работа №18. Деплой Python-приложения в Kubernetes" in block
    assert "Виды самостоятельной работы слушателей (СРС):" in block
    assert "Форма промежуточной аттестации: зачёт." in block

    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Тема 1.1. Введение в Python и установка среды разработки" in paragraph_text
    assert "Практическая работа №1. Установка Python и запуск первой программы" in paragraph_text
    assert "Практическая работа №18. Деплой Python-приложения в Kubernetes" in paragraph_text
    assert "Форма промежуточной аттестации: зачёт." in paragraph_text
    assert "Пример тестовых вопросов:" in paragraph_text
    assert "Что выведет выражение `type({1, 2, 3})`?" in paragraph_text


def test_devops_working_programs_have_blank_lines_between_topics(client, devops_seed_payload):
    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))

    working_index = _find_paragraph_index(document, "2.4. Рабочие учебные программы дисциплин/модулей")
    target_index = next(
        index
        for index in range(working_index + 1, len(document.paragraphs))
        if document.paragraphs[index].text.strip() == "Тема 1.2. Переменные и типы данных"
    )

    assert document.paragraphs[target_index - 1].text.strip() == ""


def test_working_program_formatting_uses_mixed_bold_in_section_2_4(client, devops_seed_payload):
    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))

    section_start = next(index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip().startswith("2.4."))
    section_end = next(
        index
        for index in range(section_start + 1, len(document.paragraphs))
        if document.paragraphs[index].text.strip().startswith("2.5.")
    )
    working_paragraphs = document.paragraphs[section_start + 1 : section_end]

    control_paragraph = next(
        paragraph for paragraph in working_paragraphs if paragraph.text.strip().startswith("Форма текущего контроля:")
    )
    control_runs = [run for run in control_paragraph.runs if run.text.strip()]
    assert control_runs[0].text == "Форма текущего контроля:"
    assert control_runs[0].bold is True
    assert any(run.bold is False for run in control_runs[1:])

    questions_title = next(
        paragraph for paragraph in working_paragraphs if paragraph.text.strip() == "Пример тестовых вопросов:"
    )
    question_title_runs = [run for run in questions_title.runs if run.text.strip()]
    assert question_title_runs
    assert all(run.bold is True for run in question_title_runs)

    first_question = next(paragraph for paragraph in working_paragraphs if paragraph.text.strip().startswith("1. "))
    question_runs = [run for run in first_question.runs if run.text.strip()]
    assert question_runs
    assert all(run.bold in (False, None) for run in question_runs)

    answer_paragraph = next(paragraph for paragraph in working_paragraphs if paragraph.text.strip().startswith("Ответ:"))
    answer_runs = [run for run in answer_paragraph.runs if run.text.strip()]
    assert answer_runs
    assert answer_runs[0].bold is True
    assert any(run.bold in (False, None) for run in answer_runs[1:])


def test_devops_organizational_tables_are_detailed(client, devops_seed_payload):
    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    draft = generated["draft"]

    assert len(draft["facilities"]) == 7
    assert any("Moodle" in item["equipment"] for item in draft["facilities"])
    assert any("Docker" in item["equipment"] for item in draft["facilities"])
    assert any("GitHub" in item["equipment"] for item in draft["facilities"])

    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))
    facilities_table = document.tables[9]
    facilities_text = "\n".join(cell.text for row in facilities_table.rows for cell in row.cells)

    assert "а) Материально-технические условия" in "\n".join(paragraph.text for paragraph in document.paragraphs)
    assert "Moodle" in facilities_text
    assert "Docker Engine" in facilities_text
    assert "GitHub" in facilities_text
    assert "Kubernetes" not in facilities_text


def test_export_docx(client, seed_payload):
    draft = client.post("/api/course-drafts/generate", json=seed_payload).json()
    draft_id = draft["draft_id"]
    response = client.post(f"/api/course-drafts/{draft_id}/export-docx")
    assert response.status_code == 200
    path = Path(response.json()["document_path"])
    assert path.exists()
    with zipfile.ZipFile(path) as archive:
        text = archive.read("word/document.xml").decode("utf-8", errors="ignore")
    assert "{{" not in text


def test_standard_profile_text_is_preserved_in_draft_and_document(client, seed_payload):
    generated = client.post("/api/course-drafts/generate", json=seed_payload).json()
    draft = generated["draft"]

    assert "09.02.07" in draft["general_characteristics"]["program_goal"]
    assert "Инженерия промптов для команд" in draft["general_characteristics"]["program_goal"]
    assert "09.12.2016 № 1547" in draft["general_characteristics"]["standards_basis"]
    assert "программист" in draft["general_characteristics"]["final_attestation_result"]

    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    path = Path(response.json()["document_path"])
    document = Document(path)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "ФГОС СПО 09.02.07 «Информационные системы и программирование»" in paragraph_text
    assert "09.12.2016 № 1547" in paragraph_text
    assert "программист" in paragraph_text
    assert "Инженерия промптов для команд" in paragraph_text
    assert "Трудовые функции:" in document.tables[1].rows[1].cells[1].text


def test_calendar_variants_follow_pp_load_profile(client, seed_payload):
    body = client.post("/api/course-drafts/generate", json=seed_payload).json()
    descriptions = [variant["description"] for variant in body["draft"]["calendar_variants"]]
    assert len(descriptions) == 5
    assert "3 академических часа в неделю" in descriptions[0]
    assert "включая 2 академических часа взаимодействия с преподавателем" in descriptions[0]
    assert "30 академических часов в неделю" in descriptions[4]
    assert "72 академических часа" in descriptions[0]


def test_calendar_variants_follow_pk_load_profile(client, seed_payload):
    payload = dict(seed_payload)
    payload["program_type"] = "Программа повышения квалификации"
    body = client.post("/api/course-drafts/generate", json=payload).json()
    descriptions = [variant["description"] for variant in body["draft"]["calendar_variants"]]
    assert len(descriptions) == 5
    assert descriptions[0].startswith("Вариант 1")
    assert "3 академических часа в неделю" in descriptions[0]
    assert "12 академических часов в неделю" in descriptions[2]
    assert "самостоятельной работы" in descriptions[0]


def test_calendar_variants_follow_dop_load_profile(client, seed_payload):
    payload = dict(seed_payload)
    payload["program_type"] = "Дополнительная общеобразовательная программа"
    body = client.post("/api/course-drafts/generate", json=payload).json()
    descriptions = [variant["description"] for variant in body["draft"]["calendar_variants"]]
    assert len(descriptions) == 5
    assert "(1 акад. час в неделю)" in descriptions[0]
    assert "(10 акад. часов в неделю)" in descriptions[4]
    assert "самостоятельной работы" not in descriptions[0]


def test_export_docx_layout_regression(client, seed_payload):
    generated = client.post("/api/course-drafts/generate", json=seed_payload).json()
    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    assert response.status_code == 200

    path = Path(response.json()["document_path"])
    document = Document(path)
    template = Document(Path(__file__).resolve().parents[1] / "правильный пример.docx")

    assert len(document.sections) == len(template.sections)
    assert len(document.tables) == len(template.tables)

    for generated_section, template_section in zip(document.sections, template.sections):
        assert generated_section.page_width == template_section.page_width
        assert generated_section.page_height == template_section.page_height
        assert generated_section.left_margin == template_section.left_margin
        assert generated_section.right_margin == template_section.right_margin
        assert generated_section.top_margin == template_section.top_margin
        assert generated_section.bottom_margin == template_section.bottom_margin

    assert _paragraph_signature(document.paragraphs[0]) == _paragraph_signature(template.paragraphs[0])
    assert _paragraph_signature(document.paragraphs[23]) == _paragraph_signature(template.paragraphs[23])
    assert _paragraph_signature(document.paragraphs[49]) == _paragraph_signature(template.paragraphs[49])

    working_generated = _find_paragraph_index(document, "2.4. Рабочие учебные программы дисциплин/модулей")
    working_template = _find_paragraph_index(template, "2.4. Рабочие учебные программы дисциплин/модулей")
    assert document.paragraphs[working_generated + 1].text.startswith("Модуль 1.")
    assert document.paragraphs[working_generated + 2].text.startswith("Цель:")
    assert _paragraph_signature(document.paragraphs[working_generated + 1]) == _paragraph_signature(
        template.paragraphs[working_template + 1]
    )

    assert len(document.tables[0].rows) == 1 + len(generated["draft"]["labor_functions"])
    assert len(document.tables[1].rows) == 1 + len(generated["draft"]["activity_matrix"])
    assert len(document.tables[2].rows) == 3 + len(generated["draft"]["study_plan"])
    assert len(document.tables[0].columns) == len(template.tables[0].columns)
    assert len(document.tables[2].columns) == len(template.tables[2].columns)
    assert _paragraph_signature(document.tables[0].rows[1].cells[0].paragraphs[0]) == _paragraph_signature(
        template.tables[0].rows[1].cells[0].paragraphs[0]
    )
    assert _paragraph_signature(document.tables[2].rows[3].cells[0].paragraphs[0]) == _paragraph_signature(
        template.tables[2].rows[3].cells[0].paragraphs[0]
    )
    assert _paragraph_signature(document.tables[9].rows[1].cells[0].paragraphs[0]) == _paragraph_signature(
        template.tables[9].rows[1].cells[0].paragraphs[0]
    )


def test_update_then_export(client, seed_payload):
    draft = client.post("/api/course-drafts/generate", json=seed_payload).json()
    draft_id = draft["draft_id"]
    payload = {
        "updates": {
            "working_programs_block": "Модуль 1. Обновленный блок.\nЦель: Уточнить содержание программы."
        }
    }
    update = client.put(f"/api/course-drafts/{draft_id}", json=payload)
    assert update.status_code == 200
    export = client.post(f"/api/course-drafts/{draft_id}/export-docx")
    assert export.status_code == 200


def test_confirm_saves_program_and_returns_document(client, seed_payload):
    draft = client.post("/api/course-drafts/generate", json=seed_payload).json()
    draft_id = draft["draft_id"]
    response = client.post(f"/api/course-drafts/{draft_id}/confirm")
    assert response.status_code == 200
    body = response.json()
    assert body["program_id"] == 101
    assert Path(body["document_path"]).exists()


def test_devops_working_programs_match_course_outline(client, devops_seed_payload):
    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    draft = generated["draft"]

    block = draft["working_programs_block"]
    assert "РўРµРјР° 1.1. Р’РІРµРґРµРЅРёРµ РІ Python Рё СѓСЃС‚Р°РЅРѕРІРєР° СЃСЂРµРґС‹ СЂР°Р·СЂР°Р±РѕС‚РєРё" in block
    assert "РџСЂР°РєС‚РёС‡РµСЃРєР°СЏ СЂР°Р±РѕС‚Р° в„–1. РЈСЃС‚Р°РЅРѕРІРєР° Python Рё Р·Р°РїСѓСЃРє РїРµСЂРІРѕР№ РїСЂРѕРіСЂР°РјРјС‹" in block
    assert "РџСЂР°РєС‚РёС‡РµСЃРєР°СЏ СЂР°Р±РѕС‚Р° в„–18. РџРѕРёСЃРє Рё Р·Р°РјРµРЅР° РґР°РЅРЅС‹С… СЃ РёСЃРїРѕР»СЊР·РѕРІР°РЅРёРµРј СЂРµРіСѓР»СЏСЂРЅС‹С… РІС‹СЂР°Р¶РµРЅРёР№." in block
    assert "РџСЂР°РєС‚РёС‡РµСЃРєР°СЏ СЂР°Р±РѕС‚Р° в„–18. Р”РµРїР»РѕР№ Python-РїСЂРёР»РѕР¶РµРЅРёСЏ РІ Kubernetes" in block
    assert "Р’РёРґС‹ СЃР°РјРѕСЃС‚РѕСЏС‚РµР»СЊРЅРѕР№ СЂР°Р±РѕС‚С‹ СЃР»СѓС€Р°С‚РµР»РµР№ (РЎР РЎ):" in block
    assert "Р¤РѕСЂРјР° РїСЂРѕРјРµР¶СѓС‚РѕС‡РЅРѕР№ Р°С‚С‚РµСЃС‚Р°С†РёРё:" in block

    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "РўРµРјР° 1.1. Р’РІРµРґРµРЅРёРµ РІ Python Рё СѓСЃС‚Р°РЅРѕРІРєР° СЃСЂРµРґС‹ СЂР°Р·СЂР°Р±РѕС‚РєРё" in paragraph_text
    assert "РџСЂР°РєС‚РёС‡РµСЃРєР°СЏ СЂР°Р±РѕС‚Р° в„–1. РЈСЃС‚Р°РЅРѕРІРєР° Python Рё Р·Р°РїСѓСЃРє РїРµСЂРІРѕР№ РїСЂРѕРіСЂР°РјРјС‹" in paragraph_text
    assert "РџСЂР°РєС‚РёС‡РµСЃРєР°СЏ СЂР°Р±РѕС‚Р° в„–18. Р”РµРїР»РѕР№ Python-РїСЂРёР»РѕР¶РµРЅРёСЏ РІ Kubernetes" in paragraph_text
    assert "Р¤РѕСЂРјР° РїСЂРѕРјРµР¶СѓС‚РѕС‡РЅРѕР№ Р°С‚С‚РµСЃС‚Р°С†РёРё:" in paragraph_text
    assert "РџСЂРёРјРµСЂ С‚РµСЃС‚РѕРІС‹С… РІРѕРїСЂРѕСЃРѕРІ:" in paragraph_text
    assert "Р§С‚Рѕ РІС‹РІРµРґРµС‚ РІС‹СЂР°Р¶РµРЅРёРµ `type({1, 2, 3})`?" in paragraph_text


def test_working_program_formatting_uses_mixed_bold_in_section_2_4(client, devops_seed_payload):
    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))

    section_start = next(index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip().startswith("2.4."))
    section_end = next(
        index
        for index in range(section_start + 1, len(document.paragraphs))
        if document.paragraphs[index].text.strip().startswith("2.5.")
    )
    working_paragraphs = document.paragraphs[section_start + 1 : section_end]

    control_paragraph = next(
        paragraph for paragraph in working_paragraphs if paragraph.text.strip().startswith("Р¤РѕСЂРјР° С‚РµРєСѓС‰РµРіРѕ РєРѕРЅС‚СЂРѕР»СЏ:")
    )
    control_runs = [run for run in control_paragraph.runs if run.text.strip()]
    assert control_runs[0].text == "Р¤РѕСЂРјР° С‚РµРєСѓС‰РµРіРѕ РєРѕРЅС‚СЂРѕР»СЏ:"
    assert control_runs[0].bold is True
    assert any(run.bold is False for run in control_runs[1:])

    questions_title = next(
        paragraph for paragraph in working_paragraphs if paragraph.text.strip() == "РџСЂРёРјРµСЂ С‚РµСЃС‚РѕРІС‹С… РІРѕРїСЂРѕСЃРѕРІ:"
    )
    question_title_runs = [run for run in questions_title.runs if run.text.strip()]
    assert question_title_runs
    assert all(run.bold is True for run in question_title_runs)

    first_question = next(paragraph for paragraph in working_paragraphs if paragraph.text.strip().startswith("1. "))
    question_runs = [run for run in first_question.runs if run.text.strip()]
    assert question_runs
    assert all(run.bold in (False, None) for run in question_runs)
    assert "\nA)" in first_question.text
    assert "\nРћС‚РІРµС‚:" in first_question.text
    assert any(run.text == "\nРћС‚РІРµС‚:" and run.bold is True for run in first_question.runs)


def test_devops_working_programs_match_course_outline(client, devops_seed_payload):
    topic_1 = "\u0422\u0435\u043c\u0430 1.1. \u0412\u0432\u0435\u0434\u0435\u043d\u0438\u0435 \u0432 Python \u0438 \u0443\u0441\u0442\u0430\u043d\u043e\u0432\u043a\u0430 \u0441\u0440\u0435\u0434\u044b \u0440\u0430\u0437\u0440\u0430\u0431\u043e\u0442\u043a\u0438"
    practice_1 = "\u041f\u0440\u0430\u043a\u0442\u0438\u0447\u0435\u0441\u043a\u0430\u044f \u0440\u0430\u0431\u043e\u0442\u0430 \u21161. \u0423\u0441\u0442\u0430\u043d\u043e\u0432\u043a\u0430 Python \u0438 \u0437\u0430\u043f\u0443\u0441\u043a \u043f\u0435\u0440\u0432\u043e\u0439 \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u044b"
    regex_practice = "\u041f\u0440\u0430\u043a\u0442\u0438\u0447\u0435\u0441\u043a\u0430\u044f \u0440\u0430\u0431\u043e\u0442\u0430 \u211618. \u041f\u043e\u0438\u0441\u043a \u0438 \u0437\u0430\u043c\u0435\u043d\u0430 \u0434\u0430\u043d\u043d\u044b\u0445 \u0441 \u0438\u0441\u043f\u043e\u043b\u044c\u0437\u043e\u0432\u0430\u043d\u0438\u0435\u043c \u0440\u0435\u0433\u0443\u043b\u044f\u0440\u043d\u044b\u0445 \u0432\u044b\u0440\u0430\u0436\u0435\u043d\u0438\u0439."
    kube_practice = "\u041f\u0440\u0430\u043a\u0442\u0438\u0447\u0435\u0441\u043a\u0430\u044f \u0440\u0430\u0431\u043e\u0442\u0430 \u211618. \u0414\u0435\u043f\u043b\u043e\u0439 Python-\u043f\u0440\u0438\u043b\u043e\u0436\u0435\u043d\u0438\u044f \u0432 Kubernetes"
    srs_label = "\u0412\u0438\u0434\u044b \u0441\u0430\u043c\u043e\u0441\u0442\u043e\u044f\u0442\u0435\u043b\u044c\u043d\u043e\u0439 \u0440\u0430\u0431\u043e\u0442\u044b \u0441\u043b\u0443\u0448\u0430\u0442\u0435\u043b\u0435\u0439 (\u0421\u0420\u0421):"
    attestation_label = "\u0424\u043e\u0440\u043c\u0430 \u043f\u0440\u043e\u043c\u0435\u0436\u0443\u0442\u043e\u0447\u043d\u043e\u0439 \u0430\u0442\u0442\u0435\u0441\u0442\u0430\u0446\u0438\u0438:"
    questions_label = "\u041f\u0440\u0438\u043c\u0435\u0440 \u0442\u0435\u0441\u0442\u043e\u0432\u044b\u0445 \u0432\u043e\u043f\u0440\u043e\u0441\u043e\u0432:"
    type_question = "\u0427\u0442\u043e \u0432\u044b\u0432\u0435\u0434\u0435\u0442 \u0432\u044b\u0440\u0430\u0436\u0435\u043d\u0438\u0435 `type({1, 2, 3})`?"

    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    draft = generated["draft"]

    block = draft["working_programs_block"]
    assert topic_1 in block
    assert practice_1 in block
    assert regex_practice in block
    assert kube_practice in block
    assert srs_label in block
    assert attestation_label in block

    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert topic_1 in paragraph_text
    assert practice_1 in paragraph_text
    assert kube_practice in paragraph_text
    assert attestation_label in paragraph_text
    assert questions_label in paragraph_text
    assert type_question in paragraph_text


def test_working_program_formatting_uses_mixed_bold_in_section_2_4(client, devops_seed_payload):
    control_prefix = "\u0424\u043e\u0440\u043c\u0430 \u0442\u0435\u043a\u0443\u0449\u0435\u0433\u043e \u043a\u043e\u043d\u0442\u0440\u043e\u043b\u044f:"
    questions_label = "\u041f\u0440\u0438\u043c\u0435\u0440 \u0442\u0435\u0441\u0442\u043e\u0432\u044b\u0445 \u0432\u043e\u043f\u0440\u043e\u0441\u043e\u0432:"
    answer_label = "\n\u041e\u0442\u0432\u0435\u0442:"

    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))

    section_start = next(index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip().startswith("2.4."))
    section_end = next(
        index
        for index in range(section_start + 1, len(document.paragraphs))
        if document.paragraphs[index].text.strip().startswith("2.5.")
    )
    working_paragraphs = document.paragraphs[section_start + 1 : section_end]

    control_paragraph = next(paragraph for paragraph in working_paragraphs if paragraph.text.strip().startswith(control_prefix))
    control_runs = [run for run in control_paragraph.runs if run.text.strip()]
    assert control_runs[0].text == control_prefix
    assert control_runs[0].bold is True
    assert any(run.bold is False for run in control_runs[1:])

    questions_title = next(paragraph for paragraph in working_paragraphs if paragraph.text.strip() == questions_label)
    question_title_runs = [run for run in questions_title.runs if run.text.strip()]
    assert question_title_runs
    assert all(run.bold is True for run in question_title_runs)

    first_question = next(paragraph for paragraph in working_paragraphs if paragraph.text.strip().startswith("1. "))
    question_runs = [run for run in first_question.runs if run.text.strip()]
    assert question_runs
    assert all(run.bold in (False, None) for run in question_runs if run.text != answer_label)
    assert "\nA)" in first_question.text
    assert answer_label in first_question.text
    assert any(run.text == answer_label and run.bold is True for run in first_question.runs)


def test_devops_organizational_tables_are_detailed(client, devops_seed_payload):
    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    draft = generated["draft"]
    block = draft["organizational_conditions_block"]

    assert len(draft["facilities"]) >= 11
    assert any("Moodle" in item["equipment"] for item in draft["facilities"])
    assert any("Kubernetes" in item["equipment"] for item in draft["facilities"])
    assert any("Prometheus" in item["equipment"] for item in draft["facilities"])
    assert "\u041e\u0444\u0438\u0446\u0438\u0430\u043b\u044c\u043d\u0430\u044f \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u0430\u0446\u0438\u044f:" in block
    assert "https://docs.python.org/3/" in block
    assert "https://docs.docker.com/" in block
    assert "https://znanium.ru/" in block
    assert "\u041a\u043e\u043d\u0435\u0447\u043d\u043e" not in block

    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))
    facilities_table = document.tables[9]
    digital_table = document.tables[10]
    facilities_text = "\n".join(cell.text for row in facilities_table.rows for cell in row.cells)
    digital_text = "\n".join(cell.text for row in digital_table.rows for cell in row.cells)
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert facilities_table.rows[0].cells[0].text.startswith("\u041d\u0430\u0438\u043c\u0435\u043d\u043e\u0432\u0430\u043d\u0438\u0435")
    assert len(facilities_table.rows) >= 13
    assert len(digital_table.rows) >= 8
    assert "\u0430) \u041c\u0430\u0442\u0435\u0440\u0438\u0430\u043b\u044c\u043d\u043e-\u0442\u0435\u0445\u043d\u0438\u0447\u0435\u0441\u043a\u0438\u0435 \u0443\u0441\u043b\u043e\u0432\u0438\u044f" in paragraph_text
    assert "\u0431) \u0423\u0447\u0435\u0431\u043d\u043e-\u043c\u0435\u0442\u043e\u0434\u0438\u0447\u0435\u0441\u043a\u043e\u0435 \u0438 \u0438\u043d\u0444\u043e\u0440\u043c\u0430\u0446\u0438\u043e\u043d\u043d\u043e\u0435 \u043e\u0431\u0435\u0441\u043f\u0435\u0447\u0435\u043d\u0438\u0435" in paragraph_text
    assert "\u0432) \u041a\u0430\u0434\u0440\u043e\u0432\u044b\u0435 \u0443\u0441\u043b\u043e\u0432\u0438\u044f" in paragraph_text
    assert "\u0433) \u0423\u0441\u043b\u043e\u0432\u0438\u044f \u0434\u043b\u044f \u0444\u0443\u043d\u043a\u0446\u0438\u043e\u043d\u0438\u0440\u043e\u0432\u0430\u043d\u0438\u044f \u044d\u043b\u0435\u043a\u0442\u0440\u043e\u043d\u043d\u043e\u0439 \u0438\u043d\u0444\u043e\u0440\u043c\u0430\u0446\u0438\u043e\u043d\u043d\u043e-\u043e\u0431\u0440\u0430\u0437\u043e\u0432\u0430\u0442\u0435\u043b\u044c\u043d\u043e\u0439 \u0441\u0440\u0435\u0434\u044b" in paragraph_text
    assert "\u0422\u0440\u0435\u0431\u043e\u0432\u0430\u043d\u0438\u044f \u043a \u0442\u0435\u0445\u043d\u0438\u0447\u0435\u0441\u043a\u043e\u043c\u0443 \u043e\u0431\u0435\u0441\u043f\u0435\u0447\u0435\u043d\u0438\u044e \u0441\u043b\u0443\u0448\u0430\u0442\u0435\u043b\u044f:" in paragraph_text
    assert "\u041e\u0444\u0438\u0446\u0438\u0430\u043b\u044c\u043d\u0430\u044f \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u0430\u0446\u0438\u044f:" in paragraph_text
    assert "\u042d\u043b\u0435\u043a\u0442\u0440\u043e\u043d\u043d\u044b\u0435 \u0443\u0447\u0435\u0431\u043d\u044b\u0435 \u043f\u043e\u0441\u043e\u0431\u0438\u044f:" in paragraph_text
    assert "\u041e\u0441\u043d\u043e\u0432\u043d\u044b\u0435 \u0444\u0443\u043d\u043a\u0446\u0438\u043e\u043d\u0430\u043b\u044c\u043d\u044b\u0435 \u0432\u043e\u0437\u043c\u043e\u0436\u043d\u043e\u0441\u0442\u0438 \u043f\u043b\u0430\u0442\u0444\u043e\u0440\u043c\u044b Moodle \u0438 \u0435\u0435 \u0430\u043d\u0430\u043b\u043e\u0433\u043e\u0432:" in paragraph_text
    assert "Moodle" in facilities_text
    assert "Kubernetes" in facilities_text
    assert "Prometheus" in facilities_text
    assert "Moodle" in digital_text
    assert "GitHub" in digital_text


def test_organizational_section_preserves_reference_order(client, devops_seed_payload):
    from docx.document import Document as _Document
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    def iter_block_items(parent):
        parent_elm = parent.element.body if isinstance(parent, _Document) else parent._tc
        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))
    blocks = list(iter_block_items(document))

    labels = []
    for block in blocks:
        if hasattr(block, "text"):
            text = block.text.strip()
            if text.startswith("2.5.") or text.startswith("\u0430)") or text.startswith("\u0431)") or text.startswith("\u0432)") or text.startswith("\u0433)") or text.startswith("\u041e\u0426\u0415\u041d\u041a\u0410"):
                labels.append(text)
        else:
            labels.append("TABLE")

    joined = "\n".join(labels)
    expected_sequence = [
        "2.5. \u041e\u0440\u0433\u0430\u043d\u0438\u0437\u0430\u0446\u0438\u043e\u043d\u043d\u043e-\u043f\u0435\u0434\u0430\u0433\u043e\u0433\u0438\u0447\u0435\u0441\u043a\u0438\u0435 \u0443\u0441\u043b\u043e\u0432\u0438\u044f \u0440\u0435\u0430\u043b\u0438\u0437\u0430\u0446\u0438\u0438 \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u044b",
        "\u0430) \u041c\u0430\u0442\u0435\u0440\u0438\u0430\u043b\u044c\u043d\u043e-\u0442\u0435\u0445\u043d\u0438\u0447\u0435\u0441\u043a\u0438\u0435 \u0443\u0441\u043b\u043e\u0432\u0438\u044f",
        "TABLE",
        "\u0431) \u0423\u0447\u0435\u0431\u043d\u043e-\u043c\u0435\u0442\u043e\u0434\u0438\u0447\u0435\u0441\u043a\u043e\u0435 \u0438 \u0438\u043d\u0444\u043e\u0440\u043c\u0430\u0446\u0438\u043e\u043d\u043d\u043e\u0435 \u043e\u0431\u0435\u0441\u043f\u0435\u0447\u0435\u043d\u0438\u0435",
        "\u0432) \u041a\u0430\u0434\u0440\u043e\u0432\u044b\u0435 \u0443\u0441\u043b\u043e\u0432\u0438\u044f",
        "\u0433) \u0423\u0441\u043b\u043e\u0432\u0438\u044f \u0434\u043b\u044f \u0444\u0443\u043d\u043a\u0446\u0438\u043e\u043d\u0438\u0440\u043e\u0432\u0430\u043d\u0438\u044f \u044d\u043b\u0435\u043a\u0442\u0440\u043e\u043d\u043d\u043e\u0439 \u0438\u043d\u0444\u043e\u0440\u043c\u0430\u0446\u0438\u043e\u043d\u043d\u043e-\u043e\u0431\u0440\u0430\u0437\u043e\u0432\u0430\u0442\u0435\u043b\u044c\u043d\u043e\u0439 \u0441\u0440\u0435\u0434\u044b",
        "TABLE",
        "\u041e\u0426\u0415\u041d\u041a\u0410 \u041a\u0410\u0427\u0415\u0421\u0422\u0412\u0410 \u041e\u0421\u0412\u041e\u0415\u041d\u0418\u042f \u041f\u0420\u041e\u0413\u0420\u0410\u041c\u041c\u042b",
    ]
    cursor = 0
    for item in expected_sequence:
        cursor = joined.index(item, cursor) + len(item)


def test_activity_matrix_includes_all_profile_competencies(client, seed_payload):
    generated = client.post("/api/course-drafts/generate", json=seed_payload).json()
    draft = generated["draft"]

    matrix_competencies = draft["activity_matrix"][0]["competencies"]
    expected_lines = [
        "\u041f\u041a 1.1 \u0424\u043e\u0440\u043c\u0438\u0440\u043e\u0432\u0430\u0442\u044c \u0430\u043b\u0433\u043e\u0440\u0438\u0442\u043c\u044b \u0440\u0430\u0437\u0440\u0430\u0431\u043e\u0442\u043a\u0438 \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u043d\u044b\u0445 \u043c\u043e\u0434\u0443\u043b\u0435\u0439 \u0432 \u0441\u043e\u043e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0438\u0438 \u0441 \u0442\u0435\u0445\u043d\u0438\u0447\u0435\u0441\u043a\u0438\u043c \u0437\u0430\u0434\u0430\u043d\u0438\u0435\u043c.",
        "\u041f\u041a 1.2 \u0420\u0430\u0437\u0440\u0430\u0431\u0430\u0442\u044b\u0432\u0430\u0442\u044c \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u043d\u044b\u0435 \u043c\u043e\u0434\u0443\u043b\u0438 \u0432 \u0441\u043e\u043e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0438\u0438 \u0441 \u0442\u0435\u0445\u043d\u0438\u0447\u0435\u0441\u043a\u0438\u043c \u0437\u0430\u0434\u0430\u043d\u0438\u0435\u043c.",
        "\u041f\u041a 1.3 \u0412\u044b\u043f\u043e\u043b\u043d\u044f\u0442\u044c \u043e\u0442\u043b\u0430\u0434\u043a\u0443 \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u043d\u044b\u0445 \u043c\u043e\u0434\u0443\u043b\u0435\u0439 \u0441 \u0438\u0441\u043f\u043e\u043b\u044c\u0437\u043e\u0432\u0430\u043d\u0438\u0435\u043c \u0441\u043f\u0435\u0446\u0438\u0430\u043b\u0438\u0437\u0438\u0440\u043e\u0432\u0430\u043d\u043d\u044b\u0445 \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u043d\u044b\u0445 \u0441\u0440\u0435\u0434\u0441\u0442\u0432.",
        "\u041f\u041a 2.1 \u0420\u0430\u0437\u0440\u0430\u0431\u0430\u0442\u044b\u0432\u0430\u0442\u044c \u0442\u0440\u0435\u0431\u043e\u0432\u0430\u043d\u0438\u044f \u043a \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u043d\u044b\u043c \u043c\u043e\u0434\u0443\u043b\u044f\u043c \u043d\u0430 \u043e\u0441\u043d\u043e\u0432\u0435 \u0430\u043d\u0430\u043b\u0438\u0437\u0430 \u043f\u0440\u043e\u0435\u043a\u0442\u043d\u043e\u0439 \u0438 \u0442\u0435\u0445\u043d\u0438\u0447\u0435\u0441\u043a\u043e\u0439 \u0434\u043e\u043a\u0443\u043c\u0435\u043d\u0442\u0430\u0446\u0438\u0438.",
        "\u041f\u041a 2.2 \u0412\u044b\u043f\u043e\u043b\u043d\u044f\u0442\u044c \u0438\u043d\u0442\u0435\u0433\u0440\u0430\u0446\u0438\u044e \u043c\u043e\u0434\u0443\u043b\u0435\u0439 \u0432 \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u043d\u043e\u0435 \u043e\u0431\u0435\u0441\u043f\u0435\u0447\u0435\u043d\u0438\u0435.",
        "\u041f\u041a 2.3 \u0412\u044b\u043f\u043e\u043b\u043d\u044f\u0442\u044c \u043e\u0442\u043b\u0430\u0434\u043a\u0443 \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u043d\u043e\u0433\u043e \u043c\u043e\u0434\u0443\u043b\u044f \u0441 \u0438\u0441\u043f\u043e\u043b\u044c\u0437\u043e\u0432\u0430\u043d\u0438\u0435\u043c \u0441\u043f\u0435\u0446\u0438\u0430\u043b\u0438\u0437\u0438\u0440\u043e\u0432\u0430\u043d\u043d\u044b\u0445 \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u043d\u044b\u0445 \u0441\u0440\u0435\u0434\u0441\u0442\u0432.",
        "\u041f\u041a 2.4 \u041e\u0441\u0443\u0449\u0435\u0441\u0442\u0432\u043b\u044f\u0442\u044c \u0440\u0430\u0437\u0440\u0430\u0431\u043e\u0442\u043a\u0443 \u0442\u0435\u0441\u0442\u043e\u0432\u044b\u0445 \u043d\u0430\u0431\u043e\u0440\u043e\u0432 \u0438 \u0442\u0435\u0441\u0442\u043e\u0432\u044b\u0445 \u0441\u0446\u0435\u043d\u0430\u0440\u0438\u0435\u0432 \u0434\u043b\u044f \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u043d\u043e\u0433\u043e \u043e\u0431\u0435\u0441\u043f\u0435\u0447\u0435\u043d\u0438\u044f.",
        "\u041f\u041a 2.5 \u041f\u0440\u043e\u0438\u0437\u0432\u043e\u0434\u0438\u0442\u044c \u0438\u043d\u0441\u043f\u0435\u043a\u0442\u0438\u0440\u043e\u0432\u0430\u043d\u0438\u0435 \u043a\u043e\u043c\u043f\u043e\u043d\u0435\u043d\u0442\u043e\u0432 \u043f\u0440\u043e\u0433\u0440\u0430\u043c\u043c\u043d\u043e\u0433\u043e \u043e\u0431\u0435\u0441\u043f\u0435\u0447\u0435\u043d\u0438\u044f \u043d\u0430 \u043f\u0440\u0435\u0434\u043c\u0435\u0442 \u0441\u043e\u043e\u0442\u0432\u0435\u0442\u0441\u0442\u0432\u0438\u044f \u0441\u0442\u0430\u043d\u0434\u0430\u0440\u0442\u0430\u043c \u043a\u043e\u0434\u0438\u0440\u043e\u0432\u0430\u043d\u0438\u044f.",
    ]
    for line in expected_lines:
        assert line in matrix_competencies

    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))
    matrix_text = "\n".join(cell.text for row in document.tables[1].rows for cell in row.cells)
    for line in expected_lines:
        assert line in matrix_text


def test_module_exam_examples_are_limited_to_two_and_match_module(client, devops_seed_payload):
    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    block = generated["draft"]["working_programs_block"]

    module_1 = block.split("Модуль 2.")[0]
    module_2 = block.split("Модуль 2.")[1].split("Модуль 3.")[0]
    module_3 = block.split("Модуль 3.")[1].split("Модуль 4.")[0]

    assert module_1.count("Ответ:") == 2
    assert "`re.findall" not in module_1
    assert "type({1, 2, 3})" in module_1
    assert "гарантированным закрытием" in module_1

    assert module_2.count("Ответ:") == 2
    assert "SSH" in module_2
    assert "Ansible" in module_2

    assert module_3.count("Ответ:") == 2
    assert "Dockerfile" in module_3
    assert "Deployment" in module_3


def test_testing_track_includes_all_four_competencies(client, seed_payload):
    payload = deepcopy(seed_payload)
    payload["course_name"] = "Автоматизация тестирования веб-приложений на Python"
    payload["qualification"] = "специалист по автоматизированному тестированию программного обеспечения"
    payload["professional_area"] = "Автоматизация тестирования программного обеспечения"
    payload["training_goal"] = "Освоение Python, pytest, Selenium, Playwright и инструментов контроля качества программного обеспечения"
    payload["brief_description"] = "Программа по автоматизации тестирования программного обеспечения с использованием Python и профильных инструментов."
    payload["constraints"]["standard_track_id"] = "testing"

    generated = client.post("/api/course-drafts/generate", json=payload).json()
    competencies = generated["draft"]["activity_matrix"][0]["competencies"]

    expected = [
        "ПК 3.1 Осуществлять ревьюирование программного кода в соответствии с технической документацией.",
        "ПК 3.2 Выполнять процесс измерения характеристик компонентов программного продукта.",
        "ПК 3.3 Производить исследование созданного программного кода с использованием специализированных программных средств с целью выявления ошибок и отклонения от алгоритма.",
        "ПК 3.4 Проводить сравнительный анализ программных продуктов и средств разработки.",
    ]
    for line in expected:
        assert line in competencies


def test_assessment_portfolio_and_commission_questions_follow_program(client, devops_seed_payload):
    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    draft = generated["draft"]
    assessment = draft["assessment_block"]

    assert "по всем 4 модулям" in assessment["final_attestation_form_and_goals_block"]
    assert "Модуль 1 (Программирование на языке Python):" in assessment["portfolio_requirements_block"]
    assert "Модуль 2 (Python для DevOps):" in assessment["portfolio_requirements_block"]
    assert "Модуль 3 (Работа с Docker и Kubernetes):" in assessment["portfolio_requirements_block"]
    assert "Модуль 4 (Автоматизация DevOps-процессов на Python):" in assessment["portfolio_requirements_block"]
    assert "Как воспроизводится инфраструктурный артефакт модуля «Работа с Docker и Kubernetes»" in assessment["commission_questions_block"]
    assert "Какие этапы CI/CD, мониторинга или управления конфигурациями вы реализовали в модуле «Автоматизация DevOps-процессов на Python»" in assessment["commission_questions_block"]
    assert "1. " not in assessment["commission_questions_block"]

    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))
    paragraph_text = "\n".join(paragraph.text for paragraph in document.paragraphs)

    assert "Модуль 1 (Программирование на языке Python):" in paragraph_text
    assert "Модуль 4 (Автоматизация DevOps-процессов на Python):" in paragraph_text
    assert "Как воспроизводится инфраструктурный артефакт модуля «Работа с Docker и Kubernetes»" in paragraph_text


def test_assessment_formatting_uses_template_like_bold_and_lists(client, devops_seed_payload):
    generated = client.post("/api/course-drafts/generate", json=devops_seed_payload).json()
    response = client.post(f"/api/course-drafts/{generated['draft_id']}/export-docx")
    document = Document(Path(response.json()["document_path"]))

    form_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "Форма: экзамен (защита портфолио).")
    form_runs = [run for run in form_paragraph.runs if run.text]
    assert form_runs[0].text == "Форма:"
    assert form_runs[0].bold is True
    assert any(run.bold in (False, None) for run in form_runs[1:])

    structure_paragraph = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "Структура репозитория:")
    assert all(run.bold is True for run in structure_paragraph.runs if run.text.strip())

    module_line = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip().startswith("Модуль 1 (Программирование на языке Python):"))
    assert module_line.style.name == "List Bullet"

    report_line = next(paragraph for paragraph in document.paragraphs if paragraph.text.strip() == "Краткий обзор портфолио (по модулям).")
    report_numpr = report_line._p.pPr.numPr
    assert report_numpr is not None
    assert report_line.text.strip() == "Краткий обзор портфолио (по модулям)."

    commission_question = next(paragraph for paragraph in document.paragraphs if "Как устроена серверная архитектура" in paragraph.text or "Как воспроизводится инфраструктурный артефакт" in paragraph.text)
    commission_numpr = commission_question._p.pPr.numPr
    assert commission_numpr is not None
    assert not commission_question.text.strip().startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9."))
