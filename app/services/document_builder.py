from __future__ import annotations

import re
from copy import deepcopy
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.shared import Pt

from app.config import Settings
from app.schemas.draft import CourseDraft


class DocumentBuilder:
    _PRIMARY_FONT = "Montserrat"

    def __init__(self, settings: Settings) -> None:
        self._settings = settings

    def template_available(self) -> bool:
        return self._settings.template_path.exists()

    def build_document(self, draft: CourseDraft) -> Path:
        if self._settings.template_path.exists():
            return self._build_document_from_template(draft)
        return self._build_document_fallback(draft)

    def _build_document_from_template(self, draft: CourseDraft) -> Path:
        document = Document(self._settings.template_path)

        self._fill_template_title(document, draft)
        self._fill_template_general(document, draft)
        self._fill_template_tables(document, draft)
        self._fill_template_calendar_intro(document, draft)
        self._fill_template_working_programs(document, draft)
        self._fill_template_organizational(document, draft)
        self._fill_template_assessment(document, draft)
        self._fill_template_signatures(document, draft)
        self._apply_font_to_document(document)

        file_name = f"{self._slugify(draft.program_card.course_name)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path = self._settings.output_dir / file_name
        document.save(path)
        return path

    def _build_document_fallback(self, draft: CourseDraft) -> Path:
        document = Document()
        style = document.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)

        self._build_title(document, draft)
        self._build_general_characteristics(document, draft)
        self._build_labor_functions_table(document, draft)
        self._build_activity_matrix_table(document, draft)
        self._build_study_plan_table(document, draft)
        self._build_working_programs_section(document, draft)
        self._build_calendar_variants_section(document, draft)
        self._build_organizational_section(document, draft)
        self._build_assessment_section(document, draft)
        self._build_resources_tables(document, draft)
        self._build_signatures(document, draft)

        file_name = f"{self._slugify(draft.program_card.course_name)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        path = self._settings.output_dir / file_name
        document.save(path)
        return path

    def _fill_template_title(self, document: Document, draft: CourseDraft) -> None:
        self._set_paragraph_text(document, "УТВЕРЖДАЮ", "УТВЕРЖДАЮ")
        self._set_paragraph_text(document, "Генеральный директор", draft.document_meta.approval_position)
        self._set_paragraph_text(document, "_______________ / Е. А. Шимбирева", f"_______________ / {draft.document_meta.approval_name}")
        self._set_paragraph_text(document, "«___» ____________ 2025 г.", draft.document_meta.approval_date)

        paragraph_0 = document.paragraphs[0]
        self._replace_paragraph_text(paragraph_0, self._format_organization_name(draft.document_meta.organization_name))
        self._set_paragraph_text(document, "ПРОГРАММА ПРОФЕССИОНАЛЬНОЙ ПЕРЕПОДГОТОВКИ", draft.program_card.program_type.upper())
        self._set_paragraph_text(
            document,
            "«РАЗРАБОТКА WEB-ПРИЛОЖЕНИЙ НА PYTHON (DJANGO)»",
            f"«{draft.program_card.course_name_upper}»",
        )
        self._set_paragraph_text(document, "Москва, 2025", f"{draft.document_meta.city}, {draft.document_meta.document_year}")

    def _fill_template_general(self, document: Document, draft: CourseDraft) -> None:
        gc = draft.general_characteristics
        self._set_paragraph_text(
            document,
            "Цель реализации программы — формирование у обучающихся профессиональных компетенций, необходимых для выполнения трудовых функций в области backend-разработки веб-приложений с использованием языка Python и фреймворка Django, в соответствии с ПС 06.035 «Разработчик WEB и мультимедийных приложений». (приказ Минтруда России от 18.01.2017 № 44н) и с учётом ФГОС СПО 09.02.09 Веб-разработка (Приказ Минпросвещения России от 21.11.2023 N 879).",
            gc.program_goal,
        )
        self._set_paragraph_text(document, "06.035 «Разработчик WEB и мультимедийных приложений».", self._ensure_period(gc.professional_area))
        self._replace_block_between_texts(
            document,
            "б) Объекты профессиональной деятельности",
            "в) Виды профессиональной деятельности",
            [self._ensure_period(item, use_semicolon=True, is_last=index == len(gc.professional_objects) - 1) for index, item in enumerate(gc.professional_objects)],
        )
        self._replace_block_between_texts(
            document,
            "в) Виды профессиональной деятельности",
            "г) Уровень квалификации в соответствии с профессиональным стандартом",
            [self._ensure_period(item) for item in gc.activity_types],
        )
        self._set_paragraph_text(document, "5 уровень квалификации.", self._ensure_period(gc.qualification_level))
        self._set_paragraph_text(
            document,
            "В результате освоения программы слушатель должен овладеть следующими трудовыми функциями (в соответствии с профессиональным стандартом) и/или профессиональными компетенциями (в соответствии с ФГОС):",
            "В результате освоения программы слушатель должен овладеть следующими трудовыми функциями (в соответствии с профессиональным стандартом) и/или профессиональными компетенциями (в соответствии с ФГОС):",
        )
        self._replace_block_between_texts(
            document,
            "К освоению программы допускаются:",
            "Дополнительные требования:",
            [self._ensure_period(item, use_semicolon=True, is_last=index == len(gc.audience_requirements) - 1) for index, item in enumerate(gc.audience_requirements)],
        )
        self._replace_block_between_texts(
            document,
            "Дополнительные требования:",
            "1.5. Объем программы",
            [*map(self._ensure_period, gc.additional_requirements), self._ensure_period(gc.entry_requirements)],
        )
        self._set_paragraph_text(document, "256 академических часов.", self._academic_hours_total_phrase(draft.program_card.hours))
        self._set_paragraph_text(
            document,
            "Заочная с применением электронного обучения и (или) дистанционных образовательных технологий.",
            self._ensure_period(gc.education_form),
        )
        self._set_paragraph_text(
            document,
            "В результате освоения программы и успешной сдачи итоговой аттестации слушателю выдается диплом о профессиональной переподготовке (документ о квалификации) с указанием квалификации: специалист по тестированию программного обеспечения (уровень квалификации 5 по ПС 06.035).",
            gc.final_attestation_result,
        )
        self._set_paragraph_text(
            document,
            "Лицам, осваивающим программу параллельно с получением СПО и (или) ВО, диплом о профпереподготовке выдается одновременно с получением соответствующего документа об образовании и о квалификации по базовой программе.",
            gc.parallel_education_note,
        )
        self._set_paragraph_text(
            document,
            "Программа разработана с учётом требований Профессионального стандарта 06.035 «Разработчик WEB и мультимедийных приложений». (приказ Минтруда России от 18.01.2017 № 44н), а также ФГОС СПО 09.02.09 Веб-разработка (Приказ Минпросвещения России от 21.11.2023 N 879).",
            gc.standards_basis,
        )

    def _fill_template_tables(self, document: Document, draft: CourseDraft) -> None:
        self._fill_labor_functions_table(document.tables[0], draft)
        self._fill_activity_matrix_table(document.tables[1], draft)
        self._fill_study_plan_table(document.tables[2], draft)
        for index, variant in enumerate(draft.calendar_variants, start=3):
            self._fill_calendar_variant_table(document.tables[index], variant, draft)
        self._fill_thematic_plan_table(document.tables[8], draft)
        self._apply_font_size_to_table(document.tables[8], Pt(9))
        self._fill_resources_table(document.tables[9], draft.facilities)
        self._fill_resources_table(document.tables[10], draft.digital_resources)

    def _fill_template_calendar_intro(self, document: Document, draft: CourseDraft) -> None:
        gc = draft.general_characteristics
        self._set_paragraph_text(
            document,
            "Все варианты соответствуют объему программы (256 академических часов) и обеспечивают достижение планируемых результатов. Выбор варианта осуществляется на основании заявления слушателя и фиксируется в приказе о зачислении.",
            gc.calendar_variants_intro_2,
        )
        variant_indices = [index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip().startswith("Вариант ")]
        for idx, paragraph_index in enumerate(variant_indices[: len(draft.calendar_variants)]):
            self._replace_paragraph_text(document.paragraphs[paragraph_index], draft.calendar_variants[idx].description)
        self._set_paragraph_text(
            document,
            "Учебно-тематический план программы профессиональной переподготовки",
            f"Учебно-тематический план {draft.program_card.program_type.lower()}",
        )

    def _fill_template_working_programs(self, document: Document, draft: CourseDraft) -> None:
        self._replace_block_between_texts(
            document,
            "2.4. Рабочие учебные программы дисциплин/модулей",
            "2.5. Организационно-педагогические условия реализации программы",
            self._coalesce_working_program_lines(self._split_block_preserve_empty(draft.working_programs_block)),
            paragraph_formatter=self._format_structured_working_program_paragraph,
        )

    def _fill_template_organizational(self, document: Document, draft: CourseDraft) -> None:
        start_text = "2.5. Организационно-педагогические условия реализации программы"
        end_text = "ОЦЕНКА КАЧЕСТВА ОСВОЕНИЯ ПРОГРАММЫ"
        start_index = self._find_paragraph_index(document, start_text)
        end_index = self._find_paragraph_index(document, end_text)
        existing = document.paragraphs[start_index + 1 : end_index]
        heading_candidates = [paragraph for paragraph in existing if paragraph.text.strip().startswith(("а)", "б)", "в)", "г)"))]
        heading_template = heading_candidates[0] if heading_candidates else document.paragraphs[start_index]
        body_candidates = [
            paragraph
            for paragraph in existing
            if paragraph.text.strip() and not paragraph.text.strip().startswith(("а)", "б)", "в)", "г)"))
        ]
        body_template = body_candidates[0] if body_candidates else heading_template
        for paragraph in existing:
            self._remove_paragraph(paragraph)

        sections = self._split_organizational_sections(draft.organizational_conditions_block)

        self._insert_formatted_paragraph_before_table(document, document.tables[9], heading_template, "а) Материально-технические условия")

        between_tables: list[tuple[Paragraph, str]] = []
        if sections["a"]:
            between_tables.append((body_template, ""))
            between_tables.extend((body_template, line) for line in sections["a"])
        if sections["b"]:
            between_tables.append((body_template, ""))
            between_tables.append((heading_template, "б) Учебно-методическое и информационное обеспечение"))
            between_tables.extend((body_template, line) for line in sections["b"])
        if sections["v"]:
            between_tables.append((body_template, ""))
            between_tables.append((heading_template, "в) Кадровые условия"))
            between_tables.extend((body_template, line) for line in sections["v"])
        between_tables.append((body_template, ""))
        between_tables.append((heading_template, "г) Условия для функционирования электронной информационно-образовательной среды"))
        for template, line in between_tables:
            self._insert_formatted_paragraph_before_table(document, document.tables[10], template, line)

        end_paragraph = self._find_paragraph(document, end_text)
        for line in ["", *sections["g"]]:
            paragraph = self._clone_paragraph_before(end_paragraph, body_template, line)
            self._format_organizational_paragraph(paragraph, line)

        self._replace_block_between_texts(
            document,
            "ОЦЕНКА КАЧЕСТВА ОСВОЕНИЯ ПРОГРАММЫ",
            "3.1. Форма текущего контроля",
            [
                "Оценка качества освоения программы включает:",
                "Текущий контроль – проверка выполнения практических заданий.",
                "Промежуточную аттестацию – зачет.",
                "Итоговую аттестацию – экзамен.",
            ],
        )

    def _fill_template_assessment(self, document: Document, draft: CourseDraft) -> None:
        block = draft.assessment_block
        self._replace_assessment_block(document, "3.1. Форма текущего контроля", "3.2. Форма промежуточной аттестации", self._split_block(block.current_control_block), "current_control")
        self._replace_assessment_block(document, "3.2. Форма промежуточной аттестации", "3.3. Итоговая аттестация", self._split_block(block.intermediate_attestation_block), "intermediate")
        self._replace_assessment_block(document, "3.3. Итоговая аттестация", "3.3.1. Форма и цели", self._split_block(block.final_attestation_intro_block), "final_intro")
        self._replace_assessment_block(document, "3.3.1. Форма и цели", "3.3.2. Требования к портфолио", self._split_block(block.final_attestation_form_and_goals_block), "form_and_goals")
        self._replace_assessment_block(document, "3.3.2. Требования к портфолио", "3.3.3. Порядок проведения", self._split_block(block.portfolio_requirements_block), "portfolio")
        self._replace_assessment_block(document, "3.3.3. Порядок проведения", "3.3.4. Структура доклада слушателя", self._split_block(block.attestation_procedure_block), "procedure")
        self._replace_assessment_block(document, "3.3.4. Структура доклада слушателя", "3.3.5. Примерные вопросы комиссии", self._split_block(block.report_structure_block), "report")
        self._replace_assessment_block(document, "3.3.5. Примерные вопросы комиссии", "3.3.6. Результаты и пересдача", self._split_block(block.commission_questions_block), "commission")
        self._replace_assessment_block(document, "3.3.6. Результаты и пересдача", "3.4. Критерии оценки итогового экзамена", self._split_block(block.results_and_retake_block), "results")
        self._replace_assessment_block(document, "3.4. Критерии оценки итогового экзамена", "СОСТАВИТЕЛЬИ ПРОГРАММЫ", self._split_block(block.exam_grading_criteria_block), "grading")

    def _fill_template_signatures(self, document: Document, draft: CourseDraft) -> None:
        teacher_table = document.tables[11]
        self._replace_cell_text(teacher_table.rows[0].cells[0], draft.signatures.teacher_position)
        self._replace_cell_text(teacher_table.rows[0].cells[1], draft.signatures.teacher_signature_line)
        self._replace_cell_text(teacher_table.rows[0].cells[2], draft.signatures.teacher_name)

        manager_table = document.tables[12]
        self._replace_cell_text(manager_table.rows[0].cells[0], draft.signatures.program_manager_position)
        self._replace_cell_text(manager_table.rows[0].cells[1], draft.signatures.program_manager_signature_line)
        self._replace_cell_text(manager_table.rows[0].cells[2], draft.signatures.program_manager_name)

    def _build_title(self, document: Document, draft: CourseDraft) -> None:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(draft.document_meta.organization_name.upper()).bold = True

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(draft.program_card.program_type.upper()).bold = True

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(draft.program_card.course_name_upper).bold = True

        for line in (
            f"Объём программы: {draft.program_card.hours} академических часов",
            f"Форма обучения: {draft.program_card.format}",
            f"{draft.document_meta.city}, {draft.document_meta.document_year}",
        ):
            p = document.add_paragraph(line)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _build_general_characteristics(self, document: Document, draft: CourseDraft) -> None:
        document.add_paragraph("1. Общая характеристика программы")
        gc = draft.general_characteristics
        for line in [
            f"Цель программы: {gc.program_goal}",
            f"Профессиональная область: {gc.professional_area}",
            "Объекты профессиональной деятельности:",
            *[f"- {item}" for item in gc.professional_objects],
            "Виды деятельности:",
            *[f"- {item}" for item in gc.activity_types],
            f"Требования к слушателям: {'; '.join(gc.audience_requirements)}",
            f"Дополнительные требования: {'; '.join(gc.additional_requirements)}",
            f"Итоговая аттестация: {gc.final_attestation_result}",
            gc.standards_basis,
        ]:
            self._add_non_empty_paragraph(document, line)

    def _build_labor_functions_table(self, document: Document, draft: CourseDraft) -> None:
        document.add_paragraph("2.1. Трудовые функции")
        table = document.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header = table.rows[0].cells
        header[0].text = "Наименование трудовой функции"
        header[1].text = "Код / уровень"
        header[2].text = "Компетенции"
        for item in draft.labor_functions:
            row = table.add_row().cells
            row[0].text = item.name
            row[1].text = item.code_level
            row[2].text = "\n".join(item.competencies)

    def _build_activity_matrix_table(self, document: Document, draft: CourseDraft) -> None:
        document.add_paragraph("2.2. Матрица деятельности")
        table = document.add_table(rows=1, cols=5)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        header = table.rows[0].cells
        header[0].text = "Деятельность"
        header[1].text = "Компетенции"
        header[2].text = "Практический опыт"
        header[3].text = "Умения"
        header[4].text = "Знания"
        for item in draft.activity_matrix:
            row = table.add_row().cells
            row[0].text = item.activity
            row[1].text = item.competencies
            row[2].text = item.practical_experience
            row[3].text = item.skills
            row[4].text = item.knowledge

    def _build_study_plan_table(self, document: Document, draft: CourseDraft) -> None:
        document.add_paragraph("2.3. Учебный план")
        table = document.add_table(rows=3, cols=9)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.rows[0].cells[0].text = "№"
        table.rows[0].cells[1].text = "Наименование дисциплин (модулей)"
        table.rows[0].cells[2].text = "Всего часов"
        table.rows[0].cells[3].text = "ДОТ всего"
        table.rows[0].cells[4].text = "Лекции"
        table.rows[0].cells[5].text = "Лаб."
        table.rows[0].cells[6].text = "Практика"
        table.rows[0].cells[7].text = "СРС"
        table.rows[0].cells[8].text = "Контроль"
        table.rows[1].cells[0].text = ""
        table.rows[2].cells[0].text = ""
        for item in draft.study_plan:
            row = table.add_row().cells
            row[0].text = item.number
            row[1].text = item.name
            row[2].text = str(item.total_hours)
            row[3].text = str(item.distance_total)
            row[4].text = str(item.lectures)
            row[5].text = str(item.labs)
            row[6].text = str(item.practice)
            row[7].text = str(item.srs)
            row[8].text = item.intermediate_attestation or item.current_control

    def _build_working_programs_section(self, document: Document, draft: CourseDraft) -> None:
        document.add_paragraph("2.4. Рабочие учебные программы дисциплин/модулей")
        for line in self._split_block(draft.working_programs_block):
            self._add_non_empty_paragraph(document, line)

    def _build_calendar_variants_section(self, document: Document, draft: CourseDraft) -> None:
        document.add_paragraph("2.5. Календарный учебный график")
        self._add_non_empty_paragraph(document, draft.general_characteristics.calendar_variants_intro_1)
        self._add_non_empty_paragraph(document, draft.general_characteristics.calendar_variants_intro_2)

        for variant in draft.calendar_variants:
            self._add_non_empty_paragraph(document, variant.title)
            self._add_non_empty_paragraph(document, variant.description)
            table = document.add_table(rows=1, cols=5)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            header = table.rows[0].cells
            header[0].text = "Период"
            header[1].text = "Содержание"
            header[2].text = "Часы"
            header[3].text = "С ДОТ/преп."
            header[4].text = "СРС"
            for row_data in variant.rows:
                row = table.add_row().cells
                row[0].text = row_data.period
                row[1].text = row_data.content
                row[2].text = str(row_data.total_hours)
                row[3].text = str(row_data.distance_with_teacher)
                row[4].text = str(row_data.srs)

    def _build_organizational_section(self, document: Document, draft: CourseDraft) -> None:
        document.add_paragraph("3.1. Организационно-педагогические условия")
        for line in self._split_block(draft.organizational_conditions_block):
            self._add_non_empty_paragraph(document, line)

    def _build_assessment_section(self, document: Document, draft: CourseDraft) -> None:
        block = draft.assessment_block
        section_map = [
            ("3.2. Текущий контроль", block.current_control_block),
            ("3.3. Промежуточная аттестация", block.intermediate_attestation_block),
            ("3.3.1. Итоговая аттестация", block.final_attestation_intro_block),
            ("3.3.2. Форма и цели итоговой аттестации", block.final_attestation_form_and_goals_block),
            ("3.3.3. Требования к портфолио", block.portfolio_requirements_block),
            ("3.3.4. Порядок проведения", block.attestation_procedure_block),
            ("3.3.5. Структура доклада", block.report_structure_block),
            ("3.3.6. Результаты и пересдача", block.results_and_retake_block),
            ("3.4. Критерии оценки итогового экзамена", block.exam_grading_criteria_block),
        ]
        for title, content in section_map:
            document.add_paragraph(title)
            for line in self._split_block(content):
                self._add_non_empty_paragraph(document, line)
        document.add_paragraph("Примерные вопросы комиссии")
        for line in self._split_block(block.commission_questions_block):
            self._add_non_empty_paragraph(document, line)

    def _build_resources_tables(self, document: Document, draft: CourseDraft) -> None:
        document.add_paragraph("4. Материально-техническое обеспечение")
        facilities = document.add_table(rows=1, cols=3)
        facilities.rows[0].cells[0].text = "Наименование"
        facilities.rows[0].cells[1].text = "Вид занятий"
        facilities.rows[0].cells[2].text = "Оснащение"
        for item in draft.facilities:
            row = facilities.add_row().cells
            row[0].text = item.name
            row[1].text = item.lesson_type
            row[2].text = item.equipment

        document.add_paragraph("5. Цифровые ресурсы")
        resources = document.add_table(rows=1, cols=3)
        resources.rows[0].cells[0].text = "Наименование"
        resources.rows[0].cells[1].text = "Вид занятий"
        resources.rows[0].cells[2].text = "Оснащение"
        for item in draft.digital_resources:
            row = resources.add_row().cells
            row[0].text = item.name
            row[1].text = item.lesson_type
            row[2].text = item.equipment

    def _build_signatures(self, document: Document, draft: CourseDraft) -> None:
        document.add_paragraph("Подписи")
        table = document.add_table(rows=3, cols=2)
        table.rows[0].cells[0].text = draft.document_meta.approval_position
        table.rows[0].cells[1].text = draft.document_meta.approval_name
        table.rows[1].cells[0].text = draft.signatures.teacher_position
        table.rows[1].cells[1].text = draft.signatures.teacher_name
        table.rows[2].cells[0].text = draft.signatures.program_manager_position
        table.rows[2].cells[1].text = draft.signatures.program_manager_name

    def _add_non_empty_paragraph(self, document: Document, text: str) -> None:
        clean = text.strip()
        if not clean:
            return
        paragraph = document.add_paragraph(clean)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def _split_block(self, raw: str) -> list[str]:
        return [line.strip() for line in raw.splitlines() if line.strip()]

    def _split_block_preserve_empty(self, raw: str) -> list[str]:
        lines = [line.strip() for line in raw.splitlines()]
        while lines and not lines[0]:
            lines.pop(0)
        while lines and not lines[-1]:
            lines.pop()
        return lines

    def _coalesce_working_program_lines(self, lines: list[str]) -> list[str]:
        result: list[str] = []
        index = 0

        while index < len(lines):
            line = lines[index]
            stripped = line.strip()
            if stripped != "Пример тестовых вопросов:":
                result.append(line)
                index += 1
                continue

            result.append(line)
            index += 1
            current_question: list[str] = []
            while index < len(lines):
                current = lines[index]
                current_stripped = current.strip()
                if not current_stripped:
                    if current_question:
                        result.append("\n".join(current_question))
                        current_question = []
                    result.append(current)
                    index += 1
                    break
                if current_stripped.startswith(("Модуль ", "Тема ", "Форма промежуточной аттестации:", "2.5.")):
                    if current_question:
                        result.append("\n".join(current_question))
                    break
                if re.match(r"^\d+\.", current_stripped):
                    if current_question:
                        result.append("\n".join(current_question))
                    current_question = [current_stripped]
                else:
                    if current_question:
                        current_question.append(current_stripped)
                    else:
                        result.append(current)
                index += 1
            else:
                if current_question:
                    result.append("\n".join(current_question))
            if index < len(lines) and lines[index].strip().startswith(("Модуль ", "Тема ", "Форма промежуточной аттестации:", "2.5.")):
                continue

        return result

    def _slugify(self, text: str) -> str:
        normalized = re.sub(r"[^\w\-]+", "_", text, flags=re.UNICODE)
        normalized = re.sub(r"_+", "_", normalized).strip("_")
        return normalized or "program"

    def _set_paragraph_text(self, document: Document, original_text: str, new_text: str) -> None:
        for paragraph in document.paragraphs:
            if paragraph.text.strip() == original_text:
                self._replace_paragraph_text(paragraph, new_text)
                return
        raise ValueError(f"Не найден абзац шаблона: {original_text}")

    def _replace_block_between_texts(
        self,
        document: Document,
        start_text: str,
        end_text: str,
        lines: list[str],
        paragraph_formatter=None,
    ) -> None:
        start_index = self._find_paragraph_index(document, start_text)
        end_index = self._find_paragraph_index(document, end_text)
        if end_index <= start_index:
            raise ValueError(f"Некорректные границы блока: {start_text} -> {end_text}")

        existing = document.paragraphs[start_index + 1 : end_index]
        template_paragraph = next((paragraph for paragraph in existing if paragraph.text.strip()), document.paragraphs[start_index])
        for paragraph in existing:
            self._remove_paragraph(paragraph)

        end_paragraph = self._find_paragraph(document, end_text)
        for line in lines:
            paragraph = self._clone_paragraph_before(end_paragraph, template_paragraph, line)
            if paragraph_formatter is not None:
                paragraph_formatter(paragraph, line)

    def _replace_assessment_block(
        self,
        document: Document,
        start_text: str,
        end_text: str,
        lines: list[str],
        section_id: str,
    ) -> None:
        self._replace_block_between_texts(
            document,
            start_text,
            end_text,
            lines,
            paragraph_formatter=lambda paragraph, line: self._format_assessment_paragraph(paragraph, line, section_id),
        )

    def _fill_labor_functions_table(self, table: Table, draft: CourseDraft) -> None:
        template_row = deepcopy(table.rows[1]._tr)
        self._trim_table_rows(table, 1)
        for item in draft.labor_functions:
            row = self._append_cloned_row(table, template_row).cells
            self._replace_cell_text(row[0], item.name)
            self._replace_cell_text(row[1], item.code_level)
            self._replace_cell_text(row[2], "\n".join(item.competencies))

    def _fill_activity_matrix_table(self, table: Table, draft: CourseDraft) -> None:
        template_row = deepcopy(table.rows[1]._tr)
        self._trim_table_rows(table, 1)
        for item in draft.activity_matrix:
            row = self._append_cloned_row(table, template_row).cells
            self._replace_cell_text(row[0], item.activity)
            self._replace_cell_text(row[1], item.competencies)
            self._replace_cell_text(row[2], item.practical_experience)
            self._replace_cell_text(row[3], item.skills)
            self._replace_cell_text(row[4], item.knowledge)

    def _fill_study_plan_table(self, table: Table, draft: CourseDraft) -> None:
        summary_row_template = deepcopy(table.rows[3]._tr)
        module_row_template = deepcopy(table.rows[4]._tr)
        prep_row_template = deepcopy(table.rows[-3]._tr)
        exam_row_template = deepcopy(table.rows[-2]._tr)
        total_row_template = deepcopy(table.rows[-1]._tr)
        self._trim_table_rows(table, 3)
        for item in draft.study_plan:
            if item.number == "№":
                row = self._append_cloned_row(table, summary_row_template).cells
            elif item.name == "Подготовка к итоговой аттестации":
                row = self._append_cloned_row(table, prep_row_template).cells
            elif item.name == "Проведение итоговой аттестации":
                row = self._append_cloned_row(table, exam_row_template).cells
            elif item.name == "Итого:":
                row = self._append_cloned_row(table, total_row_template).cells
            else:
                row = self._append_cloned_row(table, module_row_template).cells
            values = [
                item.number,
                item.name,
                str(item.total_hours),
                str(item.distance_total),
                str(item.lectures),
                str(item.labs),
                str(item.practice),
                str(item.srs),
                item.current_control,
                item.intermediate_attestation,
            ]
            for index, value in enumerate(values):
                self._replace_cell_text(row[index], value)

    def _fill_calendar_variant_table(self, table: Table, variant, draft: CourseDraft) -> None:
        regular_row_template = deepcopy(table.rows[1]._tr)
        attestation_row_template = deepcopy(table.rows[-2]._tr)
        total_row_template = deepcopy(table.rows[-1]._tr)
        self._trim_table_rows(table, 1)
        for row_data in variant.rows:
            template = attestation_row_template if row_data.attestation else regular_row_template
            row = self._append_cloned_row(table, template).cells
            values = [
                row_data.period,
                row_data.content,
                str(row_data.total_hours),
                str(row_data.distance_with_teacher) if row_data.distance_with_teacher else "",
                str(row_data.srs),
                str(row_data.attestation),
                str(row_data.duration_weeks),
                str(row_data.hours_per_week),
                str(row_data.teacher_hours_per_week),
            ]
            for index, value in enumerate(values):
                self._replace_cell_text(row[index], value)

        total_row = self._append_cloned_row(table, total_row_template).cells
        totals = [
            "Итого:",
            "",
            str(draft.program_card.hours),
            str(sum(row.distance_with_teacher for row in variant.rows)),
            str(sum(row.srs for row in variant.rows)),
            str(sum(row.attestation for row in variant.rows)),
            str(variant.total_weeks),
            "",
            "",
        ]
        for index, value in enumerate(totals):
            self._replace_cell_text(total_row[index], value)

    def _fill_thematic_plan_table(self, table: Table, draft: CourseDraft) -> None:
        summary_row_template = deepcopy(table.rows[3]._tr)
        module_row_template = deepcopy(table.rows[4]._tr)
        theme_row_template = deepcopy(table.rows[5]._tr)
        prep_row_template = deepcopy(table.rows[-3]._tr)
        exam_row_template = deepcopy(table.rows[-2]._tr)
        total_row_template = deepcopy(table.rows[-1]._tr)
        self._trim_table_rows(table, 3)
        module_rows = [row for row in draft.study_plan if row.number.isdigit()]
        module_total_hours = sum(row.total_hours for row in module_rows)
        module_total_distance = sum(row.distance_total for row in module_rows)
        module_total_lectures = sum(row.lectures for row in module_rows)
        module_total_labs = sum(row.labs for row in module_rows)
        module_total_practice = sum(row.practice for row in module_rows)
        module_total_srs = sum(row.srs for row in module_rows)

        summary_row = self._append_cloned_row(table, summary_row_template).cells
        summary_values = ["№ модуля", "Наименование разделов, дисциплин, тем", str(module_total_hours), str(module_total_distance), str(module_total_lectures), str(module_total_labs), str(module_total_practice), str(module_total_srs), "", ""]
        for index, value in enumerate(summary_values):
            self._replace_cell_text(summary_row[index], value)

        for module, plan_row in zip(draft.modules, module_rows):
            module_row = self._append_cloned_row(table, module_row_template).cells
            module_values = [
                f"Модуль {module.number}",
                module.name,
                str(plan_row.total_hours),
                str(plan_row.distance_total),
                str(plan_row.lectures),
                str(plan_row.labs),
                str(plan_row.practice),
                str(plan_row.srs),
                "",
                plan_row.intermediate_attestation or "Зачёт",
            ]
            for index, value in enumerate(module_values):
                self._replace_cell_text(module_row[index], value)

            for theme_number, (theme_title, theme_hours) in enumerate(self._theme_rows(module), start=1):
                distance = min(theme_hours, max(0, int(round(theme_hours * 0.67))))
                lectures = distance // 2
                practice = distance - lectures
                srs = theme_hours - distance
                row = self._append_cloned_row(table, theme_row_template).cells
                values = [
                    f"{module.number}.{theme_number}.",
                    theme_title,
                    str(theme_hours),
                    str(distance),
                    str(lectures),
                    "0",
                    str(practice),
                    str(srs),
                    "Выполнение практического задания",
                    "",
                ]
                for index, value in enumerate(values):
                    self._replace_cell_text(row[index], value)

        prep = draft.study_plan[-3]
        exam = draft.study_plan[-2]
        total = draft.study_plan[-1]

        prep_row = self._append_cloned_row(table, prep_row_template).cells
        prep_values = ["", prep.name, str(prep.total_hours), str(prep.distance_total), str(prep.lectures), str(prep.labs), str(prep.practice), str(prep.srs), "", ""]
        for index, value in enumerate(prep_values):
            self._replace_cell_text(prep_row[index], value)

        exam_row = self._append_cloned_row(table, exam_row_template).cells
        exam_values = ["", exam.name, str(exam.total_hours), str(exam.distance_total), str(exam.lectures), str(exam.labs), str(exam.practice), str(exam.srs), "", exam.intermediate_attestation]
        for index, value in enumerate(exam_values):
            self._replace_cell_text(exam_row[index], value)

        total_row = self._append_cloned_row(table, total_row_template).cells
        total_values = ["", total.name, str(total.total_hours), str(total.distance_total), str(total.lectures), str(total.labs), str(total.practice), str(total.srs), "", total.intermediate_attestation]
        for index, value in enumerate(total_values):
            self._replace_cell_text(total_row[index], value)

    def _fill_resources_table(self, table: Table, items) -> None:
        template_row = deepcopy(table.rows[1]._tr)
        self._trim_table_rows(table, 1)
        for item in items:
            row = self._append_cloned_row(table, template_row).cells
            self._replace_cell_text(row[0], item.name)
            self._replace_cell_text(row[1], item.lesson_type)
            self._replace_cell_text(row[2], item.equipment)

    def _split_organizational_sections(self, block: str) -> dict[str, list[str]]:
        sections: dict[str, list[str]] = {"a": [], "b": [], "v": [], "g": []}
        current: str | None = None
        headings = {
            "а) Материально-технические условия": "a",
            "б) Учебно-методическое и информационное обеспечение": "b",
            "в) Кадровые условия": "v",
            "г) Условия для функционирования электронной информационно-образовательной среды": "g",
        }
        for line in self._split_block_preserve_empty(block):
            normalized = line.strip()
            if normalized in headings:
                current = headings[normalized]
                continue
            if current is None:
                continue
            sections[current].append(line)
        return sections

    def _insert_formatted_paragraph_before_table(self, document: Document, table: Table, template: Paragraph, text: str) -> Paragraph:
        paragraph = self._insert_paragraph_before_table(document, table, template, text)
        self._format_organizational_paragraph(paragraph, text)
        return paragraph

    def _format_organizational_paragraph(self, paragraph: Paragraph, text: str) -> None:
        stripped = text.strip()
        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.first_line_indent = None

        if not stripped:
            self._replace_paragraph_text(paragraph, "")
            return

        bold_prefixes = (
            "а) ",
            "б) ",
            "в) ",
            "г) ",
            "Требования к техническому обеспечению слушателя:",
            "Официальная документация:",
            "Электронные учебные пособия:",
            "Квалификационные требования к преподавателям:",
            "Основные функциональные возможности платформы Moodle и ее аналогов:",
        )
        is_bold = stripped.startswith(bold_prefixes)
        self._replace_paragraph_with_segments(paragraph, [(stripped, is_bold)])

    def _trim_table_rows(self, table: Table, keep_rows: int) -> None:
        while len(table.rows) > keep_rows:
            self._remove_table_row(table.rows[-1])

    def _append_cloned_row(self, table: Table, template_row) -> any:
        table._tbl.append(deepcopy(template_row))
        return table.rows[-1]

    def _clone_paragraph_before(self, anchor: Paragraph, template: Paragraph, text: str) -> Paragraph:
        new_paragraph_element = deepcopy(template._p)
        anchor._p.addprevious(new_paragraph_element)
        paragraph = Paragraph(new_paragraph_element, anchor._parent)
        self._replace_paragraph_text(paragraph, text)
        return paragraph

    def _insert_paragraph_before_table(self, document: Document, table: Table, template: Paragraph, text: str) -> Paragraph:
        paragraph = document.add_paragraph()
        paragraph.style = template.style
        paragraph.alignment = template.alignment
        self._replace_paragraph_text(paragraph, text)
        table._tbl.addprevious(paragraph._p)
        return paragraph

    def _replace_cell_text(self, cell, text: str) -> None:
        while len(cell.paragraphs) > 1:
            self._remove_paragraph(cell.paragraphs[-1])
        paragraph = cell.paragraphs[0]
        self._replace_paragraph_text(paragraph, text)

    def _replace_paragraph_text(self, paragraph: Paragraph, new_text: str) -> None:
        original_text = paragraph.text
        leading = re.match(r"^\s*", original_text).group(0) if original_text else ""
        trailing = re.search(r"\s*$", original_text).group(0) if original_text else ""
        source_run = next((run for run in paragraph.runs if run.text), paragraph.runs[0] if paragraph.runs else None)
        for run in list(paragraph.runs):
            paragraph._p.remove(run._element)
        run = paragraph.add_run(f"{leading}{new_text.strip()}{trailing}")
        self._copy_run_format_from_run(source_run, run)

    def _format_working_program_paragraph(self, paragraph: Paragraph, text: str) -> None:
        stripped = text.strip()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if not stripped:
            self._replace_paragraph_text(paragraph, "")
            return

        is_heading = stripped.startswith(("Модуль ", "Цель:", "Тема "))
        is_label = stripped.startswith(
            (
                "Содержание:",
                "Перечень практических работ занятий:",
                "Виды самостоятельной работы слушателей (СРС):",
                "Форма текущего контроля:",
                "Форма промежуточной аттестации:",
                "Критерии оценки:",
            )
        )
        is_bullet = stripped.startswith("•")

        if is_bullet:
            paragraph.paragraph_format.left_indent = Pt(14)
            paragraph.paragraph_format.first_line_indent = Pt(0)
        else:
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = None

        for run in paragraph.runs:
            run.bold = bool(is_heading or is_label)

    def _format_structured_working_program_paragraph(self, paragraph: Paragraph, text: str) -> None:
        stripped = text.strip()

        if not stripped:
            self._replace_paragraph_text(paragraph, "")
            return

        full_bold_prefixes = (
            "Модуль ",
            "Цель:",
            "Тема ",
            "Критерии оценки:",
            "Пример тестовых вопросов:",
            "Зачет выставляется",
            "Незачет ставится",
            "Ответ:",
        )
        mixed_bold_prefixes = (
            "Содержание:",
            "Перечень практических работ занятий:",
            "Виды самостоятельной работы слушателей (СРС):",
            "Форма текущего контроля:",
            "Форма промежуточной аттестации:",
        )
        bullet_prefixes = ("вЂў", "•", "o ", "о ", "·", "-", "–")
        option_pattern = re.compile(r"^[A-DА-Г]\)")
        question_pattern = re.compile(r"^\d+\.")
        is_bullet = stripped.startswith(bullet_prefixes)

        if is_bullet:
            paragraph.paragraph_format.left_indent = Pt(14)
            paragraph.paragraph_format.first_line_indent = Pt(0)
        else:
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = None

        if stripped.startswith(("•", "вЂў", "РІР‚Сћ")):
            paragraph.paragraph_format.left_indent = Pt(28)
            paragraph.paragraph_format.first_line_indent = Pt(-10)
        elif stripped.startswith(("o ", "о ")):
            paragraph.paragraph_format.left_indent = Pt(24)
            paragraph.paragraph_format.first_line_indent = Pt(-10)

        inline_prefix = next((prefix for prefix in ("Зачет ", "Незачет ", "Ответ:") if stripped.startswith(prefix)), None)
        if inline_prefix is not None:
            suffix = stripped[len(inline_prefix) :]
            segments = [(inline_prefix, True)]
            if suffix:
                segments.append((suffix, False))
            self._replace_paragraph_with_segments(paragraph, segments)
            return

        if stripped.startswith(full_bold_prefixes):
            self._replace_paragraph_with_segments(paragraph, [(stripped, True)])
            return

        mixed_prefix = next((prefix for prefix in mixed_bold_prefixes if stripped.startswith(prefix)), None)
        if mixed_prefix is not None:
            suffix = stripped[len(mixed_prefix) :]
            segments = [(mixed_prefix, True)]
            if suffix:
                segments.append((suffix, False))
            self._replace_paragraph_with_segments(paragraph, segments)
            return

        if question_pattern.match(stripped) or option_pattern.match(stripped) or is_bullet:
            self._replace_paragraph_with_segments(paragraph, [(stripped, False)])
            return

        self._replace_paragraph_with_segments(paragraph, [(stripped, False)])

    def _replace_paragraph_with_segments(self, paragraph: Paragraph, segments: list[tuple[str, bool]]) -> None:
        source_run = next((run for run in paragraph.runs if run.text), paragraph.runs[0] if paragraph.runs else None)
        for run in list(paragraph.runs):
            paragraph._p.remove(run._element)
        for text, is_bold in segments:
            if not text:
                continue
            run = paragraph.add_run(text)
            self._copy_run_format_from_run(source_run, run)
            run.bold = is_bold

    def _format_assessment_paragraph(self, paragraph: Paragraph, text: str, section_id: str) -> None:
        stripped = text.strip()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if not stripped:
            self._replace_paragraph_text(paragraph, "")
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = None
            return

        if section_id in {"report", "commission"}:
            self._replace_paragraph_with_segments(paragraph, [(stripped, False)])
            return

        mixed_prefixes = (
            "Форма:",
            "Цели:",
            "Срок предоставления:",
            "Формат:",
            "Продолжительность защиты:",
            "Идентификация:",
            "Допустимые материалы:",
        )
        full_bold_lines = (
            "Зачет проводится в форме тестирования:",
            "Структура репозитория:",
            "Минимальный состав по модулям:",
            "Критерии оценки:",
            "Перевод баллов в отметку:",
        )

        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.first_line_indent = Pt(35.45)

        if stripped in full_bold_lines:
            self._replace_paragraph_with_segments(paragraph, [(stripped, True)])
            return

        mixed_prefix = next((prefix for prefix in mixed_prefixes if stripped.startswith(prefix)), None)
        if mixed_prefix is not None:
            suffix = stripped[len(mixed_prefix) :]
            segments = [(mixed_prefix, True)]
            if suffix:
                segments.append((suffix, False))
            if stripped.startswith("Срок предоставления:"):
                paragraph.paragraph_format.left_indent = Pt(0.05)
            self._replace_paragraph_with_segments(paragraph, segments)
            return

        if stripped.startswith(("Зачет ", "Незачет ")):
            prefix = "Зачет" if stripped.startswith("Зачет ") else "Незачет"
            suffix = stripped[len(prefix) :]
            paragraph.paragraph_format.left_indent = Pt(35.45)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            self._replace_paragraph_with_segments(paragraph, [(prefix, True), (suffix, False)])
            return

        if stripped.startswith(("«отлично»", "«хорошо»", "«удовлетворительно»", "«неудовлетворительно»")):
            mark, tail = stripped.split(" — ", 1)
            paragraph.paragraph_format.left_indent = Pt(35.45)
            paragraph.paragraph_format.first_line_indent = Pt(0)
            self._replace_paragraph_with_segments(paragraph, [(mark, True), (" — " + tail, False)])
            return

        if section_id == "portfolio" and stripped.startswith("Модуль "):
            self._set_paragraph_style_if_exists(paragraph, "List Bullet")
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.first_line_indent = Pt(14.2)
            self._replace_paragraph_with_segments(paragraph, [(stripped, False)])
            return

        self._replace_paragraph_with_segments(paragraph, [(stripped, False)])

    def _set_paragraph_style_if_exists(self, paragraph: Paragraph, style_name: str) -> None:
        try:
            paragraph.style = style_name
        except KeyError:
            return

    def _format_structured_working_program_paragraph(self, paragraph: Paragraph, text: str) -> None:
        stripped = text.strip()
        if not stripped:
            self._replace_paragraph_text(paragraph, "")
            return

        paragraph.paragraph_format.left_indent = None
        paragraph.paragraph_format.first_line_indent = None

        if "\nОтвет:" in stripped:
            before, after = stripped.rsplit("\nОтвет:", 1)
            self._replace_paragraph_with_segments(paragraph, [(before, False), ("\nОтвет:", True), (after, False)])
            return

        full_bold_prefixes = ("Модуль ", "Тема ", "Критерии оценки:", "Пример тестовых вопросов:")
        mixed_bold_prefixes = (
            "Цель:",
            "Содержание:",
            "Перечень практических работ занятий:",
            "Виды самостоятельной работы слушателей (СРС):",
            "Форма текущего контроля:",
            "Форма промежуточной аттестации:",
        )
        inline_bold_prefixes = ("Зачет", "Незачет")

        if stripped.startswith(full_bold_prefixes):
            self._replace_paragraph_with_segments(paragraph, [(stripped, True)])
            return

        mixed_prefix = next((prefix for prefix in mixed_bold_prefixes if stripped.startswith(prefix)), None)
        if mixed_prefix is not None:
            suffix = stripped[len(mixed_prefix) :]
            segments = [(mixed_prefix, True)]
            if suffix:
                segments.append((suffix, False))
            self._replace_paragraph_with_segments(paragraph, segments)
            return

        inline_prefix = next((prefix for prefix in inline_bold_prefixes if stripped.startswith(prefix)), None)
        if inline_prefix is not None:
            suffix = stripped[len(inline_prefix) :]
            segments = [(inline_prefix, True)]
            if suffix:
                segments.append((suffix, False))
            self._replace_paragraph_with_segments(paragraph, segments)
            return

        self._replace_paragraph_with_segments(paragraph, [(stripped, False)])

    def _theme_rows(self, module) -> list[tuple[str, int]]:
        themes = module.theme_titles or [module.name]
        raw = [module.hours / len(themes)] * len(themes)
        floors = [max(1, int(value)) for value in raw]
        remainder = module.hours - sum(floors)
        order = list(range(len(themes)))
        index = 0
        while remainder > 0 and order:
            floors[order[index % len(order)]] += 1
            remainder -= 1
            index += 1
        return list(zip(themes, floors))

    def _find_paragraph_index(self, document: Document, text: str) -> int:
        for index, paragraph in enumerate(document.paragraphs):
            if paragraph.text.strip() == text:
                return index
        raise ValueError(f"Не найден абзац шаблона: {text}")

    def _find_paragraph(self, document: Document, text: str) -> Paragraph:
        return document.paragraphs[self._find_paragraph_index(document, text)]

    def _remove_paragraph(self, paragraph: Paragraph) -> None:
        paragraph._element.getparent().remove(paragraph._element)

    def _remove_table_row(self, row) -> None:
        row._tr.getparent().remove(row._tr)

    def _copy_run_format_from_run(self, source_run, target_run) -> None:
        if target_run is None:
            return
        if source_run is not None:
            target_run.bold = source_run.bold
            target_run.italic = source_run.italic
            target_run.underline = source_run.underline
            target_run.font.size = source_run.font.size
        self._set_run_font_name(target_run, self._PRIMARY_FONT)

    def _set_run_font_name(self, run, font_name: str) -> None:
        if run is None:
            return
        run.font.name = font_name
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = r_pr._add_rFonts()
        r_fonts.set(qn("w:ascii"), font_name)
        r_fonts.set(qn("w:hAnsi"), font_name)
        r_fonts.set(qn("w:eastAsia"), font_name)
        r_fonts.set(qn("w:cs"), font_name)

    def _apply_font_to_document(self, document: Document) -> None:
        for paragraph in document.paragraphs:
            for run in paragraph.runs:
                self._set_run_font_name(run, self._PRIMARY_FONT)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            self._set_run_font_name(run, self._PRIMARY_FONT)

    def _apply_font_size_to_table(self, table: Table, size: Pt) -> None:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = size

    def _format_organization_name(self, organization_name: str) -> str:
        normalized = organization_name.strip()
        if normalized.upper().startswith("ООО"):
            suffix = normalized[3:].strip()
            if suffix:
                return f"ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕСТВЕННОСТЬЮ {suffix}".upper()
        return normalized.upper()

    def _ensure_period(self, text: str, use_semicolon: bool = False, is_last: bool = True) -> str:
        cleaned = text.strip().rstrip(".;")
        if use_semicolon and not is_last:
            return cleaned + ";"
        return cleaned + "."

    def _academic_hours_total_phrase(self, value: int) -> str:
        ending = self._plural(value, "академический час", "академических часа", "академических часов")
        return f"{value} {ending}."

    def _plural(self, value: int, one: str, few: str, many: str) -> str:
        remainder_100 = value % 100
        remainder_10 = value % 10
        if 11 <= remainder_100 <= 14:
            return many
        if remainder_10 == 1:
            return one
        if 2 <= remainder_10 <= 4:
            return few
        return many
